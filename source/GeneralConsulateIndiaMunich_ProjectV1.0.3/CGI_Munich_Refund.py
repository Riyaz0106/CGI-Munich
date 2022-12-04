# Developer: Shaik Riyaz
# Created for: CGI Munich

import os
import webbrowser
import comtypes.client
from docx import Document
from docx.shared import Pt
from datetime import date, datetime
from PyQt5 import QtCore, QtGui, QtWidgets

import db_utils as db_conn
import Message_DialogBox as DialogBox


class Ui_RefundWindow(object):
    def __init__(self):
        self.selectedbutton = ""
        self.misc_entry = 0
        self.icwf_entry = 0
        self.postal_entry = 0
        self.fees_entry = 0

    def setupUi(self, RefundWindow):
        self.find = None
        RefundWindow.setObjectName("RefundWindow")
        RefundWindow.resize(1020, 850)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(RefundWindow.sizePolicy().hasHeightForWidth())
        RefundWindow.setSizePolicy(sizePolicy)
        RefundWindow.setMinimumSize(QtCore.QSize(1020, 850))
        RefundWindow.setMaximumSize(QtCore.QSize(1020, 850))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        RefundWindow.setFont(font)
        self.centralwidget = QtWidgets.QWidget(RefundWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.horizontalLayoutWidget = QtWidgets.QWidget(self.centralwidget)
        self.horizontalLayoutWidget.setGeometry(QtCore.QRect(20, 20, 186, 80))
        self.horizontalLayoutWidget.setObjectName("horizontalLayoutWidget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget)
        self.horizontalLayout.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.date_lbl = QtWidgets.QLabel(self.horizontalLayoutWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(1)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.date_lbl.sizePolicy().hasHeightForWidth())
        self.date_lbl.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        self.date_lbl.setFont(font)
        self.date_lbl.setAlignment(QtCore.Qt.AlignCenter)
        self.date_lbl.setObjectName("date_lbl")
        self.horizontalLayout.addWidget(self.date_lbl)
        self.date_entry = QtWidgets.QLabel(self.horizontalLayoutWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(2)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.date_entry.sizePolicy().hasHeightForWidth())
        self.date_entry.setSizePolicy(sizePolicy)
        self.date_entry.setText("")
        self.date_entry.setAlignment(QtCore.Qt.AlignLeading|QtCore.Qt.AlignLeft|QtCore.Qt.AlignVCenter)
        self.date_entry.setObjectName("date_entry")
        self.horizontalLayout.addWidget(self.date_entry)
        self.line_6 = QtWidgets.QFrame(self.centralwidget)
        self.line_6.setGeometry(QtCore.QRect(10, 100, 1001, 20))
        self.line_6.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_6.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_6.setObjectName("line_6")
        self.email_id = QtWidgets.QLabel(self.centralwidget)
        self.email_id.setGeometry(QtCore.QRect(300, 80, 211, 21))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(11)
        self.email_id.setFont(font)
        self.email_id.setStyleSheet("color: rgb(81, 81, 81);")
        self.email_id.setAlignment(QtCore.Qt.AlignLeading|QtCore.Qt.AlignLeft|QtCore.Qt.AlignVCenter)
        self.email_id.setObjectName("email_id")
        self.horizontalLayoutWidget_2 = QtWidgets.QWidget(self.centralwidget)
        self.horizontalLayoutWidget_2.setGeometry(QtCore.QRect(710, 19, 291, 81))
        self.horizontalLayoutWidget_2.setObjectName("horizontalLayoutWidget_2")
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget_2)
        self.horizontalLayout_2.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        self.receipt_lbl = QtWidgets.QLabel(self.horizontalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        self.receipt_lbl.setFont(font)
        self.receipt_lbl.setObjectName("receipt_lbl")
        self.horizontalLayout_2.addWidget(self.receipt_lbl)
        self.rcpt_entry = QtWidgets.QLabel(self.horizontalLayoutWidget_2)
        self.rcpt_entry.setText("")
        self.rcpt_entry.setObjectName("rcpt_entry")
        self.horizontalLayout_2.addWidget(self.rcpt_entry)
        self.line_16 = QtWidgets.QFrame(self.centralwidget)
        self.line_16.setGeometry(QtCore.QRect(10, 0, 1001, 20))
        self.line_16.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_16.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_16.setObjectName("line_16")
        self.label_10 = QtWidgets.QLabel(self.centralwidget)
        self.label_10.setGeometry(QtCore.QRect(210, 10, 501, 71))
        self.label_10.setAlignment(QtCore.Qt.AlignCenter)
        self.label_10.setObjectName("label_10")
        self.line_2 = QtWidgets.QFrame(self.centralwidget)
        self.line_2.setGeometry(QtCore.QRect(0, 10, 20, 791))
        self.line_2.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_2.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_2.setObjectName("line_2")
        self.line_3 = QtWidgets.QFrame(self.centralwidget)
        self.line_3.setGeometry(QtCore.QRect(1000, 10, 20, 791))
        self.line_3.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_3.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_3.setObjectName("line_3")
        self.line_7 = QtWidgets.QFrame(self.centralwidget)
        self.line_7.setGeometry(QtCore.QRect(10, 790, 1001, 20))
        self.line_7.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_7.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_7.setObjectName("line_7")
        self.line_8 = QtWidgets.QFrame(self.centralwidget)
        self.line_8.setGeometry(QtCore.QRect(10, 270, 1001, 20))
        self.line_8.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_8.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_8.setObjectName("line_8")
        self.name_lbl = QtWidgets.QLabel(self.centralwidget)
        self.name_lbl.setGeometry(QtCore.QRect(10, 290, 210, 40))
        self.name_lbl.setAlignment(QtCore.Qt.AlignRight|QtCore.Qt.AlignTrailing|QtCore.Qt.AlignVCenter)
        self.name_lbl.setObjectName("name_lbl")
        self.nat_lbl = QtWidgets.QLabel(self.centralwidget)
        self.nat_lbl.setGeometry(QtCore.QRect(10, 180, 210, 40))
        self.nat_lbl.setAlignment(QtCore.Qt.AlignRight|QtCore.Qt.AlignTrailing|QtCore.Qt.AlignVCenter)
        self.nat_lbl.setObjectName("nat_lbl")
        self.srvc_lbl = QtWidgets.QLabel(self.centralwidget)
        self.srvc_lbl.setGeometry(QtCore.QRect(10, 340, 210, 40))
        self.srvc_lbl.setAlignment(QtCore.Qt.AlignRight|QtCore.Qt.AlignTrailing|QtCore.Qt.AlignVCenter)
        self.srvc_lbl.setObjectName("srvc_lbl")
        self.doc_lbl = QtWidgets.QLabel(self.centralwidget)
        self.doc_lbl.setGeometry(QtCore.QRect(10, 230, 210, 40))
        self.doc_lbl.setAlignment(QtCore.Qt.AlignRight|QtCore.Qt.AlignTrailing|QtCore.Qt.AlignVCenter)
        self.doc_lbl.setObjectName("doc_lbl")
        self.trans_type_lbl = QtWidgets.QLabel(self.centralwidget)
        self.trans_type_lbl.setGeometry(QtCore.QRect(10, 440, 210, 40))
        self.trans_type_lbl.setAlignment(QtCore.Qt.AlignRight|QtCore.Qt.AlignTrailing|QtCore.Qt.AlignVCenter)
        self.trans_type_lbl.setObjectName("trans_type_lbl")
        self.gratis_lbl = QtWidgets.QLabel(self.centralwidget)
        self.gratis_lbl.setGeometry(QtCore.QRect(10, 390, 210, 40))
        self.gratis_lbl.setAlignment(QtCore.Qt.AlignRight|QtCore.Qt.AlignTrailing|QtCore.Qt.AlignVCenter)
        self.gratis_lbl.setObjectName("gratis_lbl")
        self.spsrvc_lbl = QtWidgets.QLabel(self.centralwidget)
        self.spsrvc_lbl.setGeometry(QtCore.QRect(10, 490, 210, 40))
        self.spsrvc_lbl.setAlignment(QtCore.Qt.AlignRight|QtCore.Qt.AlignTrailing|QtCore.Qt.AlignVCenter)
        self.spsrvc_lbl.setObjectName("spsrvc_lbl")
        self.total_lbl = QtWidgets.QLabel(self.centralwidget)
        self.total_lbl.setGeometry(QtCore.QRect(31, 660, 51, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.total_lbl.setFont(font)
        self.total_lbl.setObjectName("total_lbl")
        self.equal = QtWidgets.QLabel(self.centralwidget)
        self.equal.setGeometry(QtCore.QRect(480, 700, 21, 16))
        self.equal.setObjectName("equal")
        self.tot_ref_entry = QtWidgets.QSpinBox(self.centralwidget)
        self.tot_ref_entry.setEnabled(False)
        self.tot_ref_entry.setGeometry(QtCore.QRect(550, 690, 171, 41))
        self.tot_ref_entry.setButtonSymbols(QtWidgets.QAbstractSpinBox.NoButtons)
        self.tot_ref_entry.setMaximum(9999)
        self.tot_ref_entry.setObjectName("tot_ref_entry")
        self.total_entry = QtWidgets.QSpinBox(self.centralwidget)
        self.total_entry.setEnabled(False)
        self.total_entry.setGeometry(QtCore.QRect(30, 690, 171, 41))
        self.total_entry.setButtonSymbols(QtWidgets.QAbstractSpinBox.NoButtons)
        self.total_entry.setMaximum(9999)
        self.total_entry.setObjectName("total_entry")
        self.tot_ref_lbl = QtWidgets.QLabel(self.centralwidget)
        self.tot_ref_lbl.setGeometry(QtCore.QRect(550, 660, 131, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.tot_ref_lbl.setFont(font)
        self.tot_ref_lbl.setObjectName("tot_ref_lbl")
        self.prnt_btn = QtWidgets.QPushButton(self.centralwidget)
        self.prnt_btn.setEnabled(False)
        self.prnt_btn.setGeometry(QtCore.QRect(20, 740, 150, 41))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.prnt_btn.sizePolicy().hasHeightForWidth())
        self.prnt_btn.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.prnt_btn.setFont(font)
        self.prnt_btn.setObjectName("prnt_btn")
        self.exit_btn = QtWidgets.QPushButton(self.centralwidget, clicked=lambda: RefundWindow.close())
        self.exit_btn.setGeometry(QtCore.QRect(850, 740, 150, 40))
        self.exit_btn.setObjectName("exit_btn")
        self.name_view = QtWidgets.QLineEdit(self.centralwidget)
        self.name_view.setEnabled(False)
        self.name_view.setGeometry(QtCore.QRect(230, 290, 581, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.name_view.setFont(font)
        self.name_view.setObjectName("name_view")
        self.nat_view = QtWidgets.QLineEdit(self.centralwidget)
        self.nat_view.setEnabled(False)
        self.nat_view.setGeometry(QtCore.QRect(230, 180, 260, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.nat_view.setFont(font)
        self.nat_view.setObjectName("nat_view")
        self.srvc_view = QtWidgets.QLineEdit(self.centralwidget)
        self.srvc_view.setEnabled(False)
        self.srvc_view.setGeometry(QtCore.QRect(230, 340, 201, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.srvc_view.setFont(font)
        self.srvc_view.setObjectName("srvc_view")
        self.docs_view = QtWidgets.QLineEdit(self.centralwidget)
        self.docs_view.setEnabled(False)
        self.docs_view.setGeometry(QtCore.QRect(230, 230, 100, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.docs_view.setFont(font)
        self.docs_view.setObjectName("docs_view")
        self.trans_type_view = QtWidgets.QLineEdit(self.centralwidget)
        self.trans_type_view.setEnabled(False)
        self.trans_type_view.setGeometry(QtCore.QRect(230, 440, 201, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.trans_type_view.setFont(font)
        self.trans_type_view.setObjectName("trans_type_view")
        self.gratis_view = QtWidgets.QLineEdit(self.centralwidget)
        self.gratis_view.setEnabled(False)
        self.gratis_view.setGeometry(QtCore.QRect(230, 390, 200, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.gratis_view.setFont(font)
        self.gratis_view.setObjectName("gratis_view")
        self.spsrvc_view = QtWidgets.QLineEdit(self.centralwidget)
        self.spsrvc_view.setEnabled(False)
        self.spsrvc_view.setGeometry(QtCore.QRect(230, 490, 300, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.spsrvc_view.setFont(font)
        self.spsrvc_view.setObjectName("spsrvc_view")
        self.srvc_desc_view = QtWidgets.QLineEdit(self.centralwidget)
        self.srvc_desc_view.setEnabled(False)
        self.srvc_desc_view.setGeometry(QtCore.QRect(430, 340, 571, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.srvc_desc_view.setFont(font)
        self.srvc_desc_view.setObjectName("srvc_desc_view")
        self.other_entry = QtWidgets.QTextEdit(self.centralwidget)
        self.other_entry.setEnabled(False)
        self.other_entry.setGeometry(QtCore.QRect(230, 540, 391, 111))
        self.other_entry.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.other_entry.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.other_entry.setObjectName("other_entry")
        self.other_lbl = QtWidgets.QLabel(self.centralwidget)
        self.other_lbl.setGeometry(QtCore.QRect(10, 540, 210, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.other_lbl.setFont(font)
        self.other_lbl.setAlignment(QtCore.Qt.AlignRight|QtCore.Qt.AlignTrailing|QtCore.Qt.AlignVCenter)
        self.other_lbl.setObjectName("other_lbl")
        self.ded_lbl = QtWidgets.QLabel(self.centralwidget)
        self.ded_lbl.setGeometry(QtCore.QRect(260, 660, 101, 31))
        self.ded_lbl.setAlignment(QtCore.Qt.AlignCenter)
        self.ded_lbl.setObjectName("ded_lbl")
        self.ded_entry = QtWidgets.QSpinBox(self.centralwidget)
        self.ded_entry.setEnabled(False)
        self.ded_entry.setGeometry(QtCore.QRect(260, 690, 171, 41))
        self.ded_entry.setButtonSymbols(QtWidgets.QAbstractSpinBox.NoButtons)
        self.ded_entry.setMaximum(9999)
        self.ded_entry.setObjectName("ded_entry")
        self.minus = QtWidgets.QLabel(self.centralwidget)
        self.minus.setGeometry(QtCore.QRect(220, 700, 21, 16))
        self.minus.setAlignment(QtCore.Qt.AlignCenter)
        self.minus.setObjectName("minus")
        self.save_btn = QtWidgets.QPushButton(self.centralwidget)
        self.save_btn.setEnabled(False)
        self.save_btn.setGeometry(QtCore.QRect(700, 740, 150, 40))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.save_btn.sizePolicy().hasHeightForWidth())
        self.save_btn.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.save_btn.setFont(font)
        self.save_btn.setObjectName("save_btn")
        self.line_5 = QtWidgets.QFrame(self.centralwidget)
        self.line_5.setGeometry(QtCore.QRect(640, 580, 20, 71))
        self.line_5.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_5.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_5.setObjectName("line_5")
        self.horizontalLayoutWidget_5 = QtWidgets.QWidget(self.centralwidget)
        self.horizontalLayoutWidget_5.setGeometry(QtCore.QRect(660, 590, 337, 51))
        self.horizontalLayoutWidget_5.setObjectName("horizontalLayoutWidget_5")
        self.post_waveicf_layout_2 = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget_5)
        self.post_waveicf_layout_2.setContentsMargins(0, 0, 0, 0)
        self.post_waveicf_layout_2.setObjectName("post_waveicf_layout_2")
        self.postExp_chk = QtWidgets.QCheckBox(self.horizontalLayoutWidget_5)
        self.postExp_chk.setEnabled(False)
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.postExp_chk.setFont(font)
        self.postExp_chk.setObjectName("postExp_chk")
        self.post_waveicf_layout_2.addWidget(self.postExp_chk)
        self.waveicf_chk = QtWidgets.QCheckBox(self.horizontalLayoutWidget_5)
        self.waveicf_chk.setEnabled(False)
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.waveicf_chk.setFont(font)
        self.waveicf_chk.setObjectName("waveicf_chk")
        self.post_waveicf_layout_2.addWidget(self.waveicf_chk)
        self.line_9 = QtWidgets.QFrame(self.centralwidget)
        self.line_9.setGeometry(QtCore.QRect(990, 580, 20, 71))
        self.line_9.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_9.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_9.setObjectName("line_9")
        self.line_17 = QtWidgets.QFrame(self.centralwidget)
        self.line_17.setGeometry(QtCore.QRect(650, 570, 351, 20))
        self.line_17.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_17.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_17.setObjectName("line_17")
        self.line_18 = QtWidgets.QFrame(self.centralwidget)
        self.line_18.setGeometry(QtCore.QRect(650, 640, 351, 20))
        self.line_18.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_18.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_18.setObjectName("line_18")
        self.edt_btn = QtWidgets.QPushButton(self.centralwidget, clicked=lambda: self.edit_Ui())
        self.edt_btn.setVisible(False)
        self.edt_btn.setGeometry(QtCore.QRect(700, 740, 150, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.edt_btn.setFont(font)
        self.edt_btn.setObjectName("edt_btn")
        self.org_rcpt_no_lbl = QtWidgets.QLabel(self.centralwidget)
        self.org_rcpt_no_lbl.setObjectName(u"org_rcpt_no_lbl")
        self.org_rcpt_no_lbl.setGeometry(QtCore.QRect(10, 120, 210, 40))
        self.org_rcpt_no_lbl.setAlignment(QtCore.Qt.AlignRight|QtCore.Qt.AlignTrailing|QtCore.Qt.AlignVCenter)
        self.org_rcpt_no_view = QtWidgets.QSpinBox(self.centralwidget)
        self.org_rcpt_no_view.setObjectName(u"org_rcpt_no_view")
        self.org_rcpt_no_view.setEnabled(False)
        self.org_rcpt_no_view.setGeometry(QtCore.QRect(230, 120, 260, 50))
        self.org_rcpt_no_view.setButtonSymbols(QtWidgets.QAbstractSpinBox.NoButtons)
        self.org_rcpt_no_view.setMaximum(1000000100)

        RefundWindow.setCentralWidget(self.centralwidget)
        self.statusbar = QtWidgets.QStatusBar(RefundWindow)
        self.statusbar.setObjectName("statusbar")
        RefundWindow.setStatusBar(self.statusbar)

        self.retranslateUi(RefundWindow)
        QtCore.QMetaObject.connectSlotsByName(RefundWindow)

        # Function Calls
        self.date_assign()
        self.total_entry.valueChanged.connect(self.recalculate_charges)
        self.ded_entry.valueChanged.connect(self.recalculate_charges)
        self.save_btn.clicked.connect(self.save_receipt)
        self.prnt_btn.clicked.connect(self.generate)

    # Date assign
    def date_assign(self):
        self.date_entry.setText(str(date.today()))

    # Receipt no. assign
    def receipt_assign(self):
        now = datetime.now()  # current date and time
        year = now.strftime("%y")
        month = now.strftime("%m")
        day = now.strftime("%d")

        index = "01"
        receipt_no = year + month + day
        selectquery = "SELECT receipt_no FROM refund_receipt " \
                      "WHERE DAY(receipt_date) LIKE DAY(CURRENT_DATE()) " \
                      "AND MONTH(receipt_date) LIKE MONTH(CURRENT_DATE()) " \
                      "AND YEAR(receipt_date) LIKE YEAR(CURRENT_DATE()) " \
                      "ORDER BY receipt_no DESC LIMIT 1 "
        records = db_conn.fetch_rows(selectquery, None)
        if records:
            for row in records:
                x = row[0][6:-1]
                index = str(int(x) + 1).zfill(len(x))
        receipt_no = receipt_no + str(index)
        self.rcpt_entry.setText(receipt_no + 'R')

    def fetchreceipt(self, records):
        msg = ""
        self.clear_Ui()

        if records:
            self.org_rcpt_no_view.setValue(records[0])
            self.name_view.setText(records[3])
            self.nat_view.setText(records[4])
            self.docs_view.setText(str(records[6]))
            self.trans_type_view.setText(records[7])
            self.spsrvc_view.setText(records[8])
            self.gratis_view.setText(records[9])
            if records[10] != 0:
                self.postExp_chk.setChecked(True)
            if records[11] != 0:
                self.waveicf_chk.setChecked(True)
            self.misc_entry = records[12]
            self.other_entry.setText(records[13])
            self.postal_entry = records[14]
            self.fees_entry = records[15]
            self.icwf_entry = records[16]
            self.total_entry.setValue(records[17])
            self.srvc_view.setText(records[20])
            self.srvc_desc_view.setText(records[19])
            if records[22] != 0:
                self.ded_entry.setValue(records[17] - records[22])
            if records[18] is not None:
                self.rcpt_entry.setText(records[18])
                self.ded_entry.setEnabled(False)
                self.save_btn.setEnabled(False)
                DialogBox.infodialog("Refund is already processed for this receipt.")
            else:
                self.receipt_assign()
                self.ded_entry.setEnabled(True)
                self.save_btn.setEnabled(True)
        else:
            msg += "Invalid Receipt Number. Please Enter a valid Receipt Number."

        return msg

    # Save data to database
    def save_receipt(self):
        if self.ded_entry.value() > self.total_entry.value():
            DialogBox.errordialog("Deduction amount cannot be greater than total amount.")
        else:
            self.selectedbutton = ""
            warningdialog = DialogBox.warningdialog("Are you sure you want to save/update this refund transaction to database?", True)
            warningdialog.buttonClicked.connect(self.msgbtn)
            warningdialog.exec_()

            if self.selectedbutton == "Yes":
                msg = self.save_data()
                if msg.find('successfully') != -1:
                    self.save_Ui()
                    DialogBox.infodialog(msg)
                else:
                    DialogBox.errordialog(msg)

    def save_data(self):
        msg = ""
        query = ""
        values = ()
        now = datetime.now()
        self.current_time = now.strftime("%H:%M:%S")
        current_date = date.today()

        find_query = "SELECT receipt_no FROM refund_receipt WHERE receipt_no LIKE %s LIMIT 1"
        records = db_conn.fetch_rows(find_query, (self.rcpt_entry.text(),))

        if not records:
            query1 = """INSERT INTO `refund_receipt` 
                        (`receipt_no`, `receipt_date`, `receipt_time`, `total_refund`) VALUES (%s,%s,%s,%s)"""
            result = db_conn.insert_rows(query1, (self.rcpt_entry.text(), current_date, self.current_time,
                                                  self.tot_ref_entry.value()))
            if result:
                query = """UPDATE receipt SET refund_receipt_no = %s WHERE receipt_no LIKE %s"""
                values = (self.rcpt_entry.text(), self.org_rcpt_no_view.value())
                msg += "Record inserted successfully."
            else:
                msg += "Failed to insert record into refund database."
        else:
            query = """UPDATE refund_receipt SET 
                        receipt_date = %s,
                        receipt_time = %s,
                        total_refund = %s
                        WHERE receipt_no LIKE %s"""
            values = (current_date, self.current_time, self.tot_ref_entry.value(), self.rcpt_entry.text())

            msg += "Record updated successfully."

        result1 = db_conn.insert_rows(query, values)
        if result1:
            return msg

    # Recalculate Refund amount
    def recalculate_charges(self):
        total = self.total_entry.value() - self.ded_entry.value()
        self.tot_ref_entry.setValue(total)

    def msgbtn(self, i):
        if i.text() == "&Yes":
            self.selectedbutton = "Yes"

    # Function to Generate PDF
    def generate(self):
        temp_dir = os.getcwd() + r'\templates\\'
        out_dir = str(os.path.expanduser('~')) + r'\Documents\ref_receipts\\'

        # Create reports folder
        if not os.path.exists(out_dir):
            os.makedirs(out_dir)

        in_file = out_dir + "receipt.docx"
        out_file = out_dir + self.name_view.text().replace(" ", "") + "_" + self.rcpt_entry.text()

        docx = Document(temp_dir+'receipt_template.docx')
        tables = docx.tables
        table = tables[0]

        total = '€' + str(self.total_entry.value())
        oth_desc = '(' + self.other_entry.toPlainText() + ')'
        fee = '€' + str(self.fees_entry)
        icwf = '€' + str(self.icwf_entry)
        misc = '€' + str(self.misc_entry)
        post = '€' + str(self.postal_entry)
        receiptNo = self.rcpt_entry.text()

        rcpt_no = table.cell(0, 0).paragraphs[0].add_run(receiptNo)
        rcpt_no.font.name = ' In black '
        rcpt_no.font.size = Pt(8)

        mdate = table.cell(1, 0).paragraphs[0].add_run(self.date_entry.text())
        mdate.font.name = ' In black '
        mdate.font.size = Pt(8)
        mtime = table.cell(1, 1).paragraphs[0].add_run(self.current_time)
        mtime.font.name = ' In black '
        mtime.font.size = Pt(8)

        services = table.cell(3, 1).paragraphs[0].add_run(self.srvc_view.text())
        services.font.name = ' In black '
        services.font.size = Pt(8)

        name_cust = table.cell(5, 1).paragraphs[0].add_run(self.name_view.text())
        name_cust.font.name = ' In black '
        name_cust.font.size = Pt(8)

        nat = table.cell(6, 1).paragraphs[0].add_run(self.nat_view.text())
        nat.font.name = ' In black '
        nat.font.size = Pt(8)

        fees = table.cell(8, 1).paragraphs[0].add_run(fee)
        fees.font.name = ' In black '
        fees.font.size = Pt(8)

        surcharge = table.cell(9, 1).paragraphs[0].add_run(icwf)
        surcharge.font.name = ' In black '
        surcharge.font.size = Pt(8)

        misc_amt = table.cell(10, 1).paragraphs[0].add_run(misc)
        misc_amt.font.name = ' In black '
        misc_amt.font.size = Pt(8)

        if self.misc_entry != 0:
            misc_desc = table.cell(11, 1).paragraphs[0].add_run(oth_desc)
            misc_desc.font.name = ' In black '
            misc_desc.font.size = Pt(8)

        pmt_mode = table.cell(12, 1).paragraphs[0].add_run(self.trans_type_view.text())
        pmt_mode.font.name = ' In black '
        pmt_mode.font.size = Pt(8)

        if self.postal_entry != 0:
            post_chrg_lbl = table.cell(13, 0).paragraphs[0].add_run('Postal Charges:')
            post_chrg_lbl.font.name = ' In black '
            post_chrg_lbl.font.size = Pt(8)
            post_amt = table.cell(13, 1).paragraphs[0].add_run(post)
            post_amt.font.name = ' In black '
            post_amt.font.size = Pt(8)

        doc_no = table.cell(14, 1).paragraphs[0].add_run(self.docs_view.text())
        doc_no.font.name = ' In black '
        doc_no.font.size = Pt(8)

        tot_amt = table.cell(16, 1).paragraphs[0].add_run(total)
        tot_amt.font.name = ' In black '
        tot_amt.font.size = Pt(8)

        ded_amt_lbl = table.cell(17, 0).paragraphs[0]
        ded_amt_lbl_run = ded_amt_lbl.add_run('Refund fees')
        ded_amt_lbl_run.font.name = ' In black '
        ded_amt_lbl_run.font.size = Pt(6)

        ded_amt = table.cell(17, 1).paragraphs[0].add_run('€' + str(self.ded_entry.value()))
        ded_amt.font.name = ' In black '
        ded_amt.font.size = Pt(8)

        tot_ref_amt_lbl = table.cell(18, 0).paragraphs[0]
        tot_ref_amt_lbl_run = tot_ref_amt_lbl.add_run(self.tot_ref_lbl.text())
        tot_ref_amt_lbl_run.font.name = ' In black '
        tot_ref_amt_lbl_run.font.size = Pt(6)

        tot_ref_amt = table.cell(18, 1).paragraphs[0].add_run('€' + str(self.tot_ref_entry.value()))
        tot_ref_amt.font.name = ' In black '
        tot_ref_amt.font.size = Pt(8)

        docx.save(in_file)

        wdFormatPDF = 17
        comtypes.CoInitialize()
        word = comtypes.client.CreateObject('Word.Application', dynamic=True)
        word.Documents.Open(in_file)
        word.Documents[0].SaveAs(out_file, wdFormatPDF)
        word.Documents[0].Close()
        word.Quit()

        if os.path.exists(in_file):
            os.remove(in_file)
            self.clear_Ui()

        returncode = webbrowser.open_new(out_file + ".pdf")
        if returncode:
            self.clear_Ui()

    def save_Ui(self):
        self.save_btn.setVisible(False)
        self.edt_btn.setVisible(True)
        self.edt_btn.setEnabled(True)
        self.ded_entry.setEnabled(False)
        self.prnt_btn.setEnabled(True)

    # Enable Edit to all data from fields
    def edit_Ui(self):
        self.ded_entry.setEnabled(True)
        self.edt_btn.setVisible(False)
        self.save_btn.setVisible(True)
        self.prnt_btn.setEnabled(False)

    # Clear all data from fields
    def clear_Ui(self):
        self.ded_entry.setEnabled(False)
        self.edt_btn.setEnabled(False)
        self.edt_btn.setVisible(False)
        self.prnt_btn.setEnabled(False)
        self.reset_Ui()

    # Set all UI elements to it's default state
    def reset_Ui(self):
        self.org_rcpt_no_view.setValue(0)
        self.rcpt_entry.setText("")
        self.name_view.setText("")
        self.docs_view.setText("")
        self.srvc_view.setText("")
        self.srvc_desc_view.setText("")
        self.nat_view.setText("")
        self.gratis_view.setText("")
        self.trans_type_view.setText("")
        self.spsrvc_view.setText("")
        self.other_entry.setText("")
        self.total_entry.setValue(0)
        self.ded_entry.setValue(0)
        self.postExp_chk.setChecked(False)
        self.waveicf_chk.setChecked(False)
        self.save_btn.setVisible(True)
        self.save_btn.setEnabled(False)

    def retranslateUi(self, RefundWindow):
        _translate = QtCore.QCoreApplication.translate
        RefundWindow.setWindowTitle(_translate("RefundWindow", "CGI Munich Refund Receipt"))
        self.date_lbl.setText(_translate("RefundWindow", "Date:"))
        self.email_id.setText(_translate("RefundWindow", "Website: cgimunich.gov.in"))
        self.receipt_lbl.setText(_translate("RefundWindow", "Receipt No:"))
        self.label_10.setText(_translate("RefundWindow", "<html><head/><body><p>"
                                                         "<img src=\"./assets/Munich_logo.png\"/>"
                                                         "</p></body></html>"))
        self.org_rcpt_no_lbl.setText(_translate("RefundWindow", "Original Receipt No:"))
        self.name_lbl.setText(_translate("RefundWindow", "Name:"))
        self.nat_lbl.setText(_translate("RefundWindow", "Nationality:"))
        self.srvc_lbl.setText(_translate("RefundWindow", "Service:"))
        self.doc_lbl.setText(_translate("RefundWindow", "No. Docs:"))
        self.trans_type_lbl.setText(_translate("RefundWindow", "Transaction Type:"))
        self.gratis_lbl.setText(_translate("RefundWindow", "Gratis:"))
        self.spsrvc_lbl.setText(_translate("RefundWindow", "Special Services:"))
        self.total_lbl.setText(_translate("RefundWindow", "Total"))
        self.equal.setText(_translate("RefundWindow", "="))
        self.tot_ref_lbl.setText(_translate("RefundWindow", "Total Refund"))
        self.prnt_btn.setText(_translate("RefundWindow", "Print"))
        self.exit_btn.setText(_translate("RefundWindow", "Exit"))
        self.other_lbl.setText(_translate("RefundWindow", "Misc Description:"))
        self.ded_lbl.setText(_translate("RefundWindow", "Deduction"))
        self.minus.setText(_translate("RefundWindow", "-"))
        self.save_btn.setText(_translate("RefundWindow", "Save"))
        self.postExp_chk.setText(_translate("RefundWindow", "Postal express"))
        self.waveicf_chk.setText(_translate("RefundWindow", "Wave ICWF"))
        self.edt_btn.setText(_translate("RefundWindow", "Edit"))


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    RefundWindow = QtWidgets.QMainWindow()
    ui = Ui_RefundWindow()
    ui.setupUi(RefundWindow)
    RefundWindow.show()
    sys.exit(app.exec_())
