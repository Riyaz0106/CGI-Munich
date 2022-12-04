# Developer: Shaik Riyaz
# Created for: CGI Munich

import os
import webbrowser
import comtypes.client
from docx import Document
from docx.shared import Pt
from datetime import date, datetime
from PyQt5 import QtCore, QtGui, QtWidgets

import db_utils as db_conn
import Message_DialogBox as DialogBox


class Ui_MainWindow(object):
    def __init__(self):
        self.selectedbutton = ""
        self.current_time = ''

    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.setEnabled(True)
        MainWindow.resize(1020, 850)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(MainWindow.sizePolicy().hasHeightForWidth())
        MainWindow.setSizePolicy(sizePolicy)
        MainWindow.setMinimumSize(QtCore.QSize(1020, 850))
        MainWindow.setMaximumSize(QtCore.QSize(1020, 850))
        font = QtGui.QFont()
        font.setPointSize(14)
        MainWindow.setFont(font)
        MainWindow.setInputMethodHints(QtCore.Qt.ImhMultiLine)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.prnt_btn = QtWidgets.QPushButton(self.centralwidget)
        self.prnt_btn.setEnabled(False)
        self.prnt_btn.setGeometry(QtCore.QRect(20, 750, 150, 41))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.prnt_btn.sizePolicy().hasHeightForWidth())
        self.prnt_btn.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.prnt_btn.setFont(font)
        self.prnt_btn.setObjectName("prnt_btn")
        self.save_btn = QtWidgets.QPushButton(self.centralwidget)
        self.save_btn.setEnabled(False)
        self.save_btn.setVisible(False)
        self.save_btn.setGeometry(QtCore.QRect(550, 750, 150, 40))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.save_btn.sizePolicy().hasHeightForWidth())
        self.save_btn.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.save_btn.setFont(font)
        self.save_btn.setObjectName("save_btn")
        self.clear_btn = QtWidgets.QPushButton(self.centralwidget)
        self.clear_btn.setEnabled(False)
        self.clear_btn.setGeometry(QtCore.QRect(700, 750, 150, 40))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.clear_btn.sizePolicy().hasHeightForWidth())
        self.clear_btn.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.clear_btn.setFont(font)
        self.clear_btn.setObjectName("clear_btn")
        self.horizontalLayoutWidget_2 = QtWidgets.QWidget(self.centralwidget)
        self.horizontalLayoutWidget_2.setGeometry(QtCore.QRect(710, 19, 291, 81))
        self.horizontalLayoutWidget_2.setObjectName("horizontalLayoutWidget_2")
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget_2)
        self.horizontalLayout_2.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        self.receipt_lbl = QtWidgets.QLabel(self.horizontalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        self.receipt_lbl.setFont(font)
        self.receipt_lbl.setObjectName("receipt_lbl")
        self.horizontalLayout_2.addWidget(self.receipt_lbl)
        self.rcpt_entry = QtWidgets.QLabel(self.horizontalLayoutWidget_2)
        self.rcpt_entry.setText("")
        self.rcpt_entry.setObjectName("rcpt_entry")
        self.horizontalLayout_2.addWidget(self.rcpt_entry)
        self.line = QtWidgets.QFrame(self.centralwidget)
        self.line.setGeometry(QtCore.QRect(10, 270, 1001, 20))
        self.line.setFrameShape(QtWidgets.QFrame.HLine)
        self.line.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line.setObjectName("line")
        self.line_8 = QtWidgets.QFrame(self.centralwidget)
        self.line_8.setGeometry(QtCore.QRect(10, 390, 20, 131))
        self.line_8.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_8.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_8.setObjectName("line_8")
        self.line_9 = QtWidgets.QFrame(self.centralwidget)
        self.line_9.setGeometry(QtCore.QRect(340, 380, 20, 141))
        self.line_9.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_9.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_9.setObjectName("line_9")
        self.line_10 = QtWidgets.QFrame(self.centralwidget)
        self.line_10.setGeometry(QtCore.QRect(20, 510, 331, 20))
        self.line_10.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_10.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_10.setObjectName("line_10")
        self.trns_lbl = QtWidgets.QLabel(self.centralwidget)
        self.trns_lbl.setGeometry(QtCore.QRect(20, 360, 211, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setBold(True)
        font.setWeight(75)
        self.trns_lbl.setFont(font)
        self.trns_lbl.setObjectName("trns_lbl")
        self.srvc_chrg_lbl = QtWidgets.QLabel(self.centralwidget)
        self.srvc_chrg_lbl.setGeometry(QtCore.QRect(650, 360, 201, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        self.srvc_chrg_lbl.setFont(font)
        self.srvc_chrg_lbl.setObjectName("srvc_chrg_lbl")
        self.line_11 = QtWidgets.QFrame(self.centralwidget)
        self.line_11.setGeometry(QtCore.QRect(230, 370, 121, 20))
        self.line_11.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_11.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_11.setObjectName("line_11")
        self.name_lbl = QtWidgets.QLabel(self.centralwidget)
        self.name_lbl.setGeometry(QtCore.QRect(10, 300, 121, 41))
        self.name_lbl.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignTrailing | QtCore.Qt.AlignVCenter)
        self.name_lbl.setObjectName("name_lbl")
        self.name_edt = QtWidgets.QTextEdit(self.centralwidget)
        self.name_edt.setEnabled(False)
        self.name_edt.setGeometry(QtCore.QRect(140, 300, 431, 40))
        self.name_edt.setInputMethodHints(QtCore.Qt.ImhMultiLine)
        self.name_edt.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.name_edt.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.name_edt.setLineWrapMode(QtWidgets.QTextEdit.NoWrap)
        self.name_edt.setObjectName("name_edt")
        self.srvc_lbl = QtWidgets.QLabel(self.centralwidget)
        self.srvc_lbl.setGeometry(QtCore.QRect(10, 180, 121, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.srvc_lbl.setFont(font)
        self.srvc_lbl.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignTrailing | QtCore.Qt.AlignVCenter)
        self.srvc_lbl.setObjectName("srvc_lbl")

        # Special Visa edit box
        self.sp_visa_entry = QtWidgets.QTextEdit(self.centralwidget)
        self.sp_visa_entry.setEnabled(False)
        self.sp_visa_entry.setGeometry(QtCore.QRect(580, 180, 416, 41))
        self.sp_visa_entry.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.sp_visa_entry.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.sp_visa_entry.setLineWrapMode(QtWidgets.QTextEdit.NoWrap)
        self.sp_visa_entry.setObjectName("sp_visa_entry")

        # Service Code Combobox
        self.srvc_cbx = QtWidgets.QComboBox(self.centralwidget)
        self.srvc_cbx.setEnabled(False)
        self.srvc_cbx.setGeometry(QtCore.QRect(360, 180, 211, 40))
        self.srvc_cbx.setEditable(True)
        self.srvc_cbx.setObjectName("srvc_cbx")

        self.doc_lbl = QtWidgets.QLabel(self.centralwidget)
        self.doc_lbl.setGeometry(QtCore.QRect(580, 120, 121, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.doc_lbl.setFont(font)
        self.doc_lbl.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignTrailing | QtCore.Qt.AlignVCenter)
        self.doc_lbl.setObjectName("doc_lbl")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(10, 120, 121, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.label_3.setFont(font)
        self.label_3.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignTrailing | QtCore.Qt.AlignVCenter)
        self.label_3.setObjectName("label_3")

        # Other Country edit box
        self.ctry_entry = QtWidgets.QTextEdit(self.centralwidget)
        self.ctry_entry.setEnabled(False)
        self.ctry_entry.setGeometry(QtCore.QRect(360, 120, 211, 41))
        self.ctry_entry.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.ctry_entry.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.ctry_entry.setLineWrapMode(QtWidgets.QTextEdit.NoWrap)
        self.ctry_entry.setObjectName("ctry_entry")

        # Nationality Combobox
        self.nat_cbx = QtWidgets.QComboBox(self.centralwidget)
        self.nat_cbx.setGeometry(QtCore.QRect(140, 120, 211, 40))
        self.nat_cbx.setEditable(True)
        self.nat_cbx.setObjectName("nat_cbx")
        self.nat_cbx.addItem("")
        self.nat_cbx.model().item(0).setEnabled(False)
        self.nat_cbx.currentIndexChanged.connect(self.enable_oth_country)
        self.enable_oth_country(self.nat_cbx.currentIndex())

        self.line_12 = QtWidgets.QFrame(self.centralwidget)
        self.line_12.setGeometry(QtCore.QRect(990, 380, 20, 141))
        self.line_12.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_12.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_12.setObjectName("line_12")
        self.line_13 = QtWidgets.QFrame(self.centralwidget)
        self.line_13.setGeometry(QtCore.QRect(640, 390, 20, 131))
        self.line_13.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_13.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_13.setObjectName("line_13")
        self.line_14 = QtWidgets.QFrame(self.centralwidget)
        self.line_14.setGeometry(QtCore.QRect(650, 510, 351, 20))
        self.line_14.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_14.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_14.setObjectName("line_14")
        self.line_15 = QtWidgets.QFrame(self.centralwidget)
        self.line_15.setGeometry(QtCore.QRect(850, 370, 151, 20))
        self.line_15.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_15.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_15.setObjectName("line_15")
        self.line_2 = QtWidgets.QFrame(self.centralwidget)
        self.line_2.setGeometry(QtCore.QRect(0, 10, 20, 791))
        self.line_2.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_2.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_2.setObjectName("line_2")
        self.line_3 = QtWidgets.QFrame(self.centralwidget)
        self.line_3.setGeometry(QtCore.QRect(1000, 10, 20, 791))
        self.line_3.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_3.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_3.setObjectName("line_3")
        self.line_4 = QtWidgets.QFrame(self.centralwidget)
        self.line_4.setGeometry(QtCore.QRect(10, 790, 1001, 20))
        self.line_4.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_4.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_4.setObjectName("line_4")
        self.line_16 = QtWidgets.QFrame(self.centralwidget)
        self.line_16.setGeometry(QtCore.QRect(10, 0, 1001, 20))
        self.line_16.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_16.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_16.setObjectName("line_16")
        self.total_lbl = QtWidgets.QLabel(self.centralwidget)
        self.total_lbl.setGeometry(QtCore.QRect(830, 630, 51, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.total_lbl.setFont(font)
        self.total_lbl.setObjectName("total_lbl")
        self.post_lbl = QtWidgets.QLabel(self.centralwidget)
        self.post_lbl.setGeometry(QtCore.QRect(210, 630, 121, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.post_lbl.setFont(font)
        self.post_lbl.setObjectName("post_lbl")
        self.misc_lbl = QtWidgets.QLabel(self.centralwidget)
        self.misc_lbl.setGeometry(QtCore.QRect(20, 630, 101, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.misc_lbl.setFont(font)
        self.misc_lbl.setObjectName("misc_lbl")
        self.icwf_lbl = QtWidgets.QLabel(self.centralwidget)
        self.icwf_lbl.setGeometry(QtCore.QRect(590, 630, 101, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.icwf_lbl.setFont(font)
        self.icwf_lbl.setObjectName("icwf_lbl")
        self.fees_lbl = QtWidgets.QLabel(self.centralwidget)
        self.fees_lbl.setGeometry(QtCore.QRect(400, 630, 101, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.fees_lbl.setFont(font)
        self.fees_lbl.setObjectName("fees_lbl")
        self.other_entry = QtWidgets.QTextEdit(self.centralwidget)
        self.other_entry.setEnabled(False)
        self.other_entry.setGeometry(QtCore.QRect(20, 570, 331, 41))
        self.other_entry.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.other_entry.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.other_entry.setObjectName("other_entry")
        self.other_lbl = QtWidgets.QLabel(self.centralwidget)
        self.other_lbl.setGeometry(QtCore.QRect(20, 530, 181, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.other_lbl.setFont(font)
        self.other_lbl.setObjectName("other_lbl")
        self.horizontalLayoutWidget = QtWidgets.QWidget(self.centralwidget)
        self.horizontalLayoutWidget.setGeometry(QtCore.QRect(20, 20, 186, 80))
        self.horizontalLayoutWidget.setObjectName("horizontalLayoutWidget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget)
        self.horizontalLayout.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.date_lbl = QtWidgets.QLabel(self.horizontalLayoutWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(1)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.date_lbl.sizePolicy().hasHeightForWidth())
        self.date_lbl.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        self.date_lbl.setFont(font)
        self.date_lbl.setAlignment(QtCore.Qt.AlignCenter)
        self.date_lbl.setObjectName("date_lbl")
        self.horizontalLayout.addWidget(self.date_lbl)
        self.date_entry = QtWidgets.QLabel(self.horizontalLayoutWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(2)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.date_entry.sizePolicy().hasHeightForWidth())
        self.date_entry.setSizePolicy(sizePolicy)
        self.date_entry.setText("")
        self.date_entry.setAlignment(QtCore.Qt.AlignLeading | QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter)
        self.date_entry.setObjectName("date_entry")
        self.horizontalLayout.addWidget(self.date_entry)

        self.line_6 = QtWidgets.QFrame(self.centralwidget)
        self.line_6.setGeometry(QtCore.QRect(10, 100, 1001, 20))
        self.line_6.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_6.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_6.setObjectName("line_6")
        self.srvc_entry = QtWidgets.QTextEdit(self.centralwidget)
        self.srvc_entry.setEnabled(False)
        self.srvc_entry.setGeometry(QtCore.QRect(140, 230, 856, 40))
        self.srvc_entry.setObjectName("srvc_entry")
        self.label_10 = QtWidgets.QLabel(self.centralwidget)
        self.label_10.setGeometry(QtCore.QRect(210, 10, 501, 71))
        self.label_10.setAlignment(QtCore.Qt.AlignCenter)
        self.label_10.setObjectName("label_10")
        self.horizontalLayoutWidget_5 = QtWidgets.QWidget(self.centralwidget)
        self.horizontalLayoutWidget_5.setGeometry(QtCore.QRect(660, 550, 337, 51))
        self.horizontalLayoutWidget_5.setObjectName("horizontalLayoutWidget_5")
        self.post_waveicwf_layout = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget_5)
        self.post_waveicwf_layout.setContentsMargins(0, 0, 0, 0)
        self.post_waveicwf_layout.setObjectName("post_waveicwf_layout")
        self.pExpress_chk = QtWidgets.QCheckBox(self.horizontalLayoutWidget_5)
        self.pExpress_chk.setEnabled(False)
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.pExpress_chk.setFont(font)
        self.pExpress_chk.setObjectName("pExpress_chk")
        self.post_waveicwf_layout.addWidget(self.pExpress_chk)
        self.wave_chk = QtWidgets.QCheckBox(self.horizontalLayoutWidget_5)
        self.wave_chk.setEnabled(False)
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.wave_chk.setFont(font)
        self.wave_chk.setObjectName("wave_chk")
        self.post_waveicwf_layout.addWidget(self.wave_chk)
        self.doc_entry = QtWidgets.QSpinBox(self.centralwidget)
        self.doc_entry.setGeometry(QtCore.QRect(710, 120, 81, 40))
        self.doc_entry.setMinimum(1)
        self.doc_entry.setMaximum(1000)
        self.doc_entry.setObjectName("doc_entry")
        self.gridLayoutWidget = QtWidgets.QWidget(self.centralwidget)
        self.gridLayoutWidget.setGeometry(QtCore.QRect(660, 390, 331, 121))
        self.gridLayoutWidget.setObjectName("gridLayoutWidget")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.gridLayoutWidget)
        self.gridLayout_2.setContentsMargins(0, 0, 0, 0)
        self.gridLayout_2.setObjectName("gridLayout_2")
        self.Rwrkday_rbtn = QtWidgets.QRadioButton(self.gridLayoutWidget)
        self.Rwrkday_rbtn.setEnabled(False)
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.Rwrkday_rbtn.setFont(font)
        self.Rwrkday_rbtn.setObjectName("Rwrkday_rbtn")
        self.gridLayout_2.addWidget(self.Rwrkday_rbtn, 1, 0, 1, 1)
        self.Rwrkday_rbtn.toggled.connect(self.set_RWday)
        self.gratis_rbtn = QtWidgets.QRadioButton(self.gridLayoutWidget)
        self.gratis_rbtn.setEnabled(False)
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.gratis_rbtn.setFont(font)
        self.gratis_rbtn.setObjectName("gratis_rbtn")
        self.gridLayout_2.addWidget(self.gratis_rbtn, 0, 0, 1, 1)
        self.gratis_rbtn.toggled.connect(self.set_Gfees)
        self.Rholiday_rbtn = QtWidgets.QRadioButton(self.gridLayoutWidget)
        self.Rholiday_rbtn.setEnabled(False)
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.Rholiday_rbtn.setFont(font)
        self.Rholiday_rbtn.setObjectName("Rholiday_rbtn")
        self.gridLayout_2.addWidget(self.Rholiday_rbtn, 1, 1, 1, 1)
        self.Rholiday_rbtn.toggled.connect(self.set_RHday)
        self.service_chrg_grp = QtWidgets.QButtonGroup(self.centralwidget)
        self.service_chrg_grp.addButton(self.Rholiday_rbtn)
        self.service_chrg_grp.addButton(self.Rwrkday_rbtn)
        self.edt_btn = QtWidgets.QPushButton(self.centralwidget, clicked=lambda: self.edit_Ui())
        self.edt_btn.setEnabled(False)
        self.edt_btn.setVisible(False)
        self.edt_btn.setGeometry(QtCore.QRect(550, 750, 150, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.edt_btn.setFont(font)
        self.edt_btn.setObjectName("edt_btn")
        self.line_5 = QtWidgets.QFrame(self.centralwidget)
        self.line_5.setGeometry(QtCore.QRect(640, 540, 20, 71))
        self.line_5.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_5.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_5.setObjectName("line_5")
        self.line_7 = QtWidgets.QFrame(self.centralwidget)
        self.line_7.setGeometry(QtCore.QRect(1000, 540, 3, 71))
        self.line_7.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_7.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_7.setObjectName("line_7")
        self.line_17 = QtWidgets.QFrame(self.centralwidget)
        self.line_17.setGeometry(QtCore.QRect(650, 530, 351, 20))
        self.line_17.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_17.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_17.setObjectName("line_17")
        self.line_18 = QtWidgets.QFrame(self.centralwidget)
        self.line_18.setGeometry(QtCore.QRect(650, 600, 351, 20))
        self.line_18.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_18.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_18.setObjectName("line_18")

        # Service Type Combobox
        self.srvt_cbx = QtWidgets.QComboBox(self.centralwidget)
        self.srvt_cbx.setGeometry(QtCore.QRect(140, 180, 211, 40))
        self.srvt_cbx.setEditable(False)
        self.srvt_cbx.setObjectName("srvt_cbx")
        self.srvt_cbx.addItem("")
        self.srvt_cbx.model().item(0).setEnabled(False)

        self.gridLayoutWidget_2 = QtWidgets.QWidget(self.centralwidget)
        self.gridLayoutWidget_2.setGeometry(QtCore.QRect(30, 400, 311, 111))
        self.gridLayoutWidget_2.setObjectName("gridLayoutWidget_2")
        self.gridLayout_3 = QtWidgets.QGridLayout(self.gridLayoutWidget_2)
        self.gridLayout_3.setContentsMargins(0, 0, 0, 0)
        self.gridLayout_3.setObjectName("gridLayout_3")
        self.card_rbtn = QtWidgets.QRadioButton(self.gridLayoutWidget_2)
        self.card_rbtn.setEnabled(False)
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.card_rbtn.setFont(font)
        self.card_rbtn.setObjectName("card_rbtn")
        self.gridLayout_3.addWidget(self.card_rbtn, 1, 1, 1, 1)
        self.cash_rbtn = QtWidgets.QRadioButton(self.gridLayoutWidget_2)
        self.cash_rbtn.setEnabled(False)
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.cash_rbtn.setFont(font)
        self.cash_rbtn.setObjectName("cash_rbtn")
        self.gridLayout_3.addWidget(self.cash_rbtn, 1, 0, 1, 1)
        self.bank_rbtn = QtWidgets.QRadioButton(self.gridLayoutWidget_2)
        self.bank_rbtn.setChecked(True)
        self.bank_rbtn.setEnabled(False)
        font = QtGui.QFont()
        font.setFamily("Arial")
        self.bank_rbtn.setFont(font)
        self.bank_rbtn.setObjectName("bank_rbtn")
        self.gridLayout_3.addWidget(self.bank_rbtn, 0, 0, 1, 1)
        self.trans_grp = QtWidgets.QButtonGroup(self.centralwidget)
        self.trans_grp.addButton(self.bank_rbtn)
        self.trans_grp.addButton(self.cash_rbtn)
        self.trans_grp.addButton(self.card_rbtn)
        self.email_id = QtWidgets.QLabel(self.centralwidget)
        self.email_id.setGeometry(QtCore.QRect(300, 80, 211, 21))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(11)
        self.email_id.setFont(font)
        self.email_id.setStyleSheet("color: rgb(81, 81, 81);")
        self.email_id.setAlignment(QtCore.Qt.AlignLeading | QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter)
        self.email_id.setObjectName("email_id")
        self.misc_entry = QtWidgets.QSpinBox(self.centralwidget)
        self.misc_entry.setEnabled(False)
        self.misc_entry.setGeometry(QtCore.QRect(20, 660, 150, 40))
        self.misc_entry.setButtonSymbols(QtWidgets.QAbstractSpinBox.NoButtons)
        self.misc_entry.setMaximum(9999)
        self.misc_entry.setObjectName("misc_entry")
        self.postal_entry = QtWidgets.QSpinBox(self.centralwidget)
        self.postal_entry.setEnabled(False)
        self.postal_entry.setGeometry(QtCore.QRect(210, 660, 150, 41))
        self.postal_entry.setButtonSymbols(QtWidgets.QAbstractSpinBox.NoButtons)
        self.postal_entry.setMaximum(9999)
        self.postal_entry.setObjectName("postal_entry")
        self.icwf_entry = QtWidgets.QSpinBox(self.centralwidget)
        self.icwf_entry.setEnabled(False)
        self.icwf_entry.setGeometry(QtCore.QRect(590, 660, 150, 41))
        self.icwf_entry.setButtonSymbols(QtWidgets.QAbstractSpinBox.NoButtons)
        self.icwf_entry.setMaximum(9999)
        self.icwf_entry.setObjectName("icwf_entry")
        self.fees_entry = QtWidgets.QSpinBox(self.centralwidget)
        self.fees_entry.setEnabled(False)
        self.fees_entry.setGeometry(QtCore.QRect(400, 660, 150, 41))
        self.fees_entry.setButtonSymbols(QtWidgets.QAbstractSpinBox.NoButtons)
        self.fees_entry.setMaximum(9999)
        self.fees_entry.setObjectName("fees_entry")
        self.total_entry = QtWidgets.QSpinBox(self.centralwidget)
        self.total_entry.setEnabled(False)
        self.total_entry.setGeometry(QtCore.QRect(829, 660, 171, 41))
        self.total_entry.setButtonSymbols(QtWidgets.QAbstractSpinBox.NoButtons)
        self.total_entry.setMaximum(9999)
        self.total_entry.setObjectName("total_entry")
        self.exit_btn = QtWidgets.QPushButton(self.centralwidget, clicked=lambda: MainWindow.close())
        self.exit_btn.setGeometry(QtCore.QRect(850, 750, 150, 40))
        self.exit_btn.setObjectName("exit_btn")
        self.plus = QtWidgets.QLabel(self.centralwidget)
        self.plus.setGeometry(QtCore.QRect(180, 670, 21, 16))
        self.plus.setObjectName("plus")
        self.plus_2 = QtWidgets.QLabel(self.centralwidget)
        self.plus_2.setGeometry(QtCore.QRect(370, 670, 21, 16))
        self.plus_2.setObjectName("plus_2")
        self.plus_3 = QtWidgets.QLabel(self.centralwidget)
        self.plus_3.setGeometry(QtCore.QRect(560, 670, 21, 16))
        self.plus_3.setObjectName("plus_3")
        self.equal = QtWidgets.QLabel(self.centralwidget)
        self.equal.setGeometry(QtCore.QRect(780, 670, 21, 16))
        self.equal.setObjectName("equal")
        MainWindow.setCentralWidget(self.centralwidget)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

        # Function calls
        self.date_assign()
        self.receipt_assign()
        self.get_country_list()
        self.get_service_types()
        self.srvt_cbx.currentIndexChanged.connect(self.get_service_codes)
        self.get_service_codes(self.srvt_cbx.currentIndex())
        self.srvc_cbx.currentIndexChanged.connect(self.enable_special_visa)
        self.enable_special_visa(self.srvc_cbx.currentIndex())
        self.srvc_cbx.currentIndexChanged.connect(self.fetch_charges)

        self.misc_entry.valueChanged.connect(self.recalculate_charges)
        self.fees_entry.valueChanged.connect(self.recalculate_charges)
        self.postal_entry.valueChanged.connect(self.recalculate_charges)
        self.icwf_entry.valueChanged.connect(self.recalculate_charges)

        self.pExpress_chk.stateChanged.connect(self.pExpress_statechange)
        self.wave_chk.stateChanged.connect(self.waveicwf_statechange)

        self.save_btn.clicked.connect(self.save_receipt)

        self.prnt_btn.clicked.connect(self.generate)
        self.clear_btn.clicked.connect(self.clear_Ui)

    # Date assign
    def date_assign(self):
        self.date_entry.setText(str(date.today()))

    # Receipt no. assign
    def receipt_assign(self):
        now = datetime.now()  # current date and time
        year = now.strftime("%y")
        month = now.strftime("%m")
        day = now.strftime("%d")

        index = "01"
        receipt_no = year + month + day
        selectquery = "SELECT receipt_no FROM receipt " \
                      "WHERE DAY(receipt_date) LIKE DAY(CURRENT_DATE()) " \
                      "AND MONTH(receipt_date) LIKE MONTH(CURRENT_DATE()) " \
                      "AND YEAR(receipt_date) LIKE YEAR(CURRENT_DATE()) " \
                      "ORDER BY receipt_no DESC LIMIT 1 "
        records = db_conn.fetch_rows(selectquery, None)
        if records:
            for row in records:
                x = str(row[0])[6:]
                index = str(int(x) + 1).zfill(len(x))
        receipt_no = receipt_no + str(index)
        self.rcpt_entry.setText(receipt_no)

    # Get Nationality
    def get_country_list(self):
        selectquery = "SELECT name FROM nationality"
        records = db_conn.fetch_rows(selectquery, None)
        if records:
            for row in records:
                self.nat_cbx.addItem(row[0])
        else:
            DialogBox.errordialog("There are no countries in nationality database.")

    # Enable other country field
    def enable_oth_country(self, code):
        if self.nat_cbx.itemText(code) == "Other":
            self.ctry_entry.setText("")
            self.ctry_entry.setEnabled(True)
        else:
            self.ctry_entry.setText("")
            self.ctry_entry.setEnabled(False)

    # Get Service Type
    def get_service_types(self):
        items = []
        items1 = []

        selectquery = "SELECT code,category,sub_category FROM services"
        records = db_conn.fetch_rows(selectquery, None)
        if records:
            for row in records:
                items.append(row[0])
                items1.append(row[1])

            k = sorted(set(items1))
            serv_dict = dict(zip(items, items1))
            for i in k:
                final_key = []
                for key in serv_dict:
                    if i == serv_dict[key]:
                        final_key.append(key)
                self.srvt_cbx.addItem(i, final_key)
        else:
            DialogBox.errordialog("There are no services available in the Services database.")

    # Generate service codes based on service type
    def get_service_codes(self, index):
        if self.srvt_cbx.currentText() != "Service Type":
            self.srvc_cbx.setEnabled(True)
            self.srvc_cbx.clear()
            codes = self.srvt_cbx.itemData(index)
            if codes:
                self.srvc_cbx.addItems(codes)
        else:
            self.srvc_cbx.clear()
            self.srvc_cbx.setEnabled(False)

    # Enable special service visa field
    def enable_special_visa(self, code):
        if self.srvc_cbx.itemText(code) == "SV000000":
            self.sp_visa_entry.setText("")
            self.sp_visa_entry.setEnabled(True)
        else:
            self.sp_visa_entry.setText("")
            self.sp_visa_entry.setEnabled(False)

    # Search button functionality
    def fetch_charges(self):
        if self.srvc_cbx.currentText() != "":
            self.reset_Ui()
            selectquery = 'SELECT * FROM services WHERE code LIKE %s'
            records = db_conn.fetch_rows(selectquery, (self.srvc_cbx.currentText(),))
            if records:
                for row in records:
                    self.srvc_entry.setText(row[1])
                    self.fees_entry.setValue(int(row[2]) * self.doc_entry.value())
                    self.icwf_entry.setValue(int(row[3]) * self.doc_entry.value())
                    self.total_entry.setValue(self.fees_entry.value() + self.icwf_entry.value())

                self.edit_Ui()
            else:
                DialogBox.errordialog("There is no data available for the code ["
                                      + self.srvc_cbx.currentText() + ']')

    # Function for Service charges
    def set_Gfees(self):
        if self.gratis_rbtn.isChecked():
            self.fees_entry.setValue(0)
            self.icwf_entry.setValue(0)

    # Function for Service charges
    def set_RWday(self):
        if self.Rwrkday_rbtn.isChecked():
            self.Rwrkday_rbtn.setAutoExclusive(False)
            self.Rwrkday_rbtn.setChecked(False)
            self.Rwrkday_rbtn.setAutoExclusive(True)

    def set_RHday(self):
        if self.Rholiday_rbtn.isChecked():
            self.Rholiday_rbtn.setAutoExclusive(False)
            self.Rholiday_rbtn.setChecked(False)
            self.Rholiday_rbtn.setAutoExclusive(True)

    # Function for postal express state change
    def pExpress_statechange(self):
        if self.pExpress_chk.isChecked():
            pExpress_amt = self.postal_entry.value() + 4
            self.postal_entry.setValue(pExpress_amt)
        else:
            if self.postal_entry.value() != 0:
                pExpress_amt = self.postal_entry.value() - 4
                self.postal_entry.setValue(pExpress_amt)

    # Function for wave icwf state change
    def waveicwf_statechange(self):
        if self.wave_chk.isChecked():
            if self.icwf_entry.value() != 0:
                self.icwf_entry.setValue(0)

    # ReCalculate the total amount if there is any change in the values
    def recalculate_charges(self):
        total = self.misc_entry.value() + self.fees_entry.value() + self.icwf_entry.value() + \
                self.postal_entry.value()
        self.total_entry.setValue(total)
        if self.fees_entry.value() == 0 and self.srvt_cbx.currentText() != "Service Type":
            self.gratis_rbtn.setAutoExclusive(False)
            self.gratis_rbtn.setChecked(True)
            self.gratis_rbtn.setAutoExclusive(True)
        else:
            self.gratis_rbtn.setAutoExclusive(False)
            self.gratis_rbtn.setChecked(False)
            self.gratis_rbtn.setAutoExclusive(True)

    # Validations
    def field_validations(self):
        msg = ""
        if self.ctry_entry.isEnabled() and self.ctry_entry.toPlainText() == "":
            msg += 'Please specify a country name or select a different country.\n'

        if self.sp_visa_entry.isEnabled() and self.sp_visa_entry.toPlainText() == "":
            msg += 'Please enter Special visa description or select a different code.\n'

        if self.name_edt.toPlainText() == "":
            msg += 'Name field cannot be blank.\n'
        elif not self.name_edt.toPlainText().replace(" ", "").isalpha():
            msg += 'Name should contain only alphabets.\n'

        if self.misc_entry.value() != 0 and self.other_entry.toPlainText() == "":
            msg += 'Misc description field cannot be blank.\n'

        if self.nat_cbx.currentText() != "":
            if self.nat_cbx.currentText() != "Select Country":
                find_nat_query = "SELECT name FROM nationality WHERE name LIKE %s LIMIT 1"
                nat_records = db_conn.fetch_rows(find_nat_query, (self.nat_cbx.currentText(),))
                if not nat_records:
                    msg += 'Please select a valid country from the list.\n'
            else:
                msg += 'Please select a country from the list.\n'
        else:
            msg += 'Please select a country from the list.\n'

        if self.srvt_cbx.currentText() == "Service Type":
            msg += 'Please select a valid service type from the list.\n'
        elif self.srvc_cbx.currentText() != "":
            find_srvc_query = "SELECT code FROM services WHERE category LIKE %s AND code LIKE %s LIMIT 1"
            value = [self.srvt_cbx.currentText(), self.srvc_cbx.currentText()]
            values = tuple(value, )
            srvc_records = db_conn.fetch_rows(find_srvc_query, values)
            if not srvc_records:
                msg += 'Please select a valid Service code from the list.\n'
        else:
            msg += 'Service code cannot be blank.\n'
        return msg

    # Save data to database and disable all fields
    def save_receipt(self):
        errors = self.field_validations()
        if errors == "":
            self.selectedbutton = ""
            warningdialog = DialogBox.warningdialog(
                "Are you sure you want to add/update this payment transaction to database?", True)
            warningdialog.buttonClicked.connect(self.msgbtn)
            warningdialog.exec_()

            if self.selectedbutton == "Yes":
                msg = self.save_data()  # Save data to database
                if msg.find('successfully') != -1:
                    self.edt_btn.setVisible(True)
                    self.edt_btn.setEnabled(True)
                    self.save_btn.setVisible(False)
                    self.save_Ui()  # Disable all edit fields
                    DialogBox.infodialog(msg)

        else:
            DialogBox.errordialog(errors)

    # Save data to database
    def save_data(self):
        msg = ""
        post_express = 0
        wave_icwf = 0
        servc_chrg_type = "None"
        self.current_time = ""
        service_code = self.srvc_cbx.currentText()

        if self.ctry_entry.toPlainText() == "":
            country = self.nat_cbx.currentText()
        else:
            country = self.ctry_entry.toPlainText()

        name = self.name_edt.toPlainText()
        no_doc = self.doc_entry.value()

        if self.Rwrkday_rbtn.isChecked():
            servc_chrg_type = self.Rwrkday_rbtn.text()
        elif self.Rholiday_rbtn.isChecked():
            servc_chrg_type = self.Rholiday_rbtn.text()

        if self.pExpress_chk.isChecked():
            post_express = 1
        if self.wave_chk.isChecked():
            wave_icwf = 1

        misc_desc = self.other_entry.toPlainText()
        misc_amt = self.misc_entry.value()
        post_amt = self.postal_entry.value()
        fees_amt = self.fees_entry.value()
        icwf_amt = self.icwf_entry.value()
        total_amt = self.total_entry.value()
        receipt_num = self.rcpt_entry.text()
        now = datetime.now()

        self.current_time = now.strftime("%H:%M:%S")

        current_date = date.today()

        find_query = "SELECT receipt_no FROM receipt WHERE receipt_no LIKE %s"
        records = db_conn.fetch_rows(find_query, (receipt_num,))

        if not records:
            query = """INSERT INTO 
                `receipt`(`receipt_no`, `receipt_date`, `receipt_time`, `Name`, `Nationality`, `service_code`, `no_doc`, 
                `trans_type`, `serv_chrg_type`, `gratis_ind`, `post_express`, `wave_icwf`, `misc_amt`, `misc_desc`, 
                `post_amt`, `fees_amt`, `icwf_amt`, `total_amt`) 
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
            values = (
                receipt_num, current_date, self.current_time, name, country, service_code, no_doc,
                self.trans_grp.checkedButton().text(), servc_chrg_type,
                self.gratis_rbtn.text() if self.gratis_rbtn.isChecked() else 'Non Gratis', post_express, wave_icwf,
                misc_amt, misc_desc, post_amt, fees_amt, icwf_amt, total_amt)

            msg += "Record inserted successfully!!!"
        else:
            query = """UPDATE receipt SET 
                        receipt_date = %s,
                        receipt_time = %s,
                        Name = %s,
                        Nationality = %s,
                        service_code = %s,
                        no_doc = %s,
                        trans_type = %s,
                        serv_chrg_type = %s,
                        gratis_ind = %s,
                        post_express = %s,
                        wave_icwf = %s,
                        misc_amt = %s,
                        misc_desc = %s,
                        post_amt = %s,
                        fees_amt = %s,
                        icwf_amt = %s,
                        total_amt = %s
                        WHERE receipt_no LIKE %s"""
            values = (current_date, self.current_time, name, country, service_code, no_doc,
                      self.trans_grp.checkedButton().text(), servc_chrg_type,
                      self.gratis_rbtn.text() if self.gratis_rbtn.isChecked() else 'Non Gratis',
                      post_express, wave_icwf, misc_amt, misc_desc, post_amt, fees_amt, icwf_amt, total_amt,
                      receipt_num)

            msg += "Record updated successfully!!!"

        result = db_conn.insert_rows(query, values)
        if result:
            return msg

    def msgbtn(self, i):
        if i.text() == "&Yes":
            self.selectedbutton = "Yes"

    # Function to Generate PDF
    def generate(self):
        temp_dir = os.getcwd() + r'\templates\\'
        out_dir = str(os.path.expanduser('~')) + r'\Documents\orig_receipts\\'

        # Create reports folder
        if not os.path.exists(out_dir):
            os.makedirs(out_dir)

        in_file = out_dir + 'receipt.docx'
        out_file = out_dir + self.name_edt.toPlainText().replace(" ", "") + "_" + self.rcpt_entry.text()

        docx = Document(temp_dir + 'receipt_template.docx')
        tables = docx.tables
        table = tables[0]

        total = '€' + str(self.total_entry.value())
        oth_desc = '(' + self.other_entry.toPlainText() + ')'
        fee = '€' + str(self.fees_entry.value())
        icwf = '€' + str(self.icwf_entry.value())
        misc = '€' + str(self.misc_entry.value())
        post = '€' + str(self.postal_entry.value())
        receiptNo = self.rcpt_entry.text()

        rcpt_no = table.cell(0, 0).paragraphs[0].add_run(receiptNo)
        rcpt_no.font.name = ' In black '
        rcpt_no.font.size = Pt(8)

        mdate = table.cell(1, 0).paragraphs[0].add_run(self.date_entry.text())
        mdate.font.name = ' In black '
        mdate.font.size = Pt(8)
        mtime = table.cell(1, 1).paragraphs[0].add_run(self.current_time)
        mtime.font.name = ' In black '
        mtime.font.size = Pt(8)

        if self.sp_visa_entry.toPlainText() == "":
            services = table.cell(3, 1).paragraphs[0].add_run(self.srvc_entry.toPlainText())
            services.font.name = ' In black '
            services.font.size = Pt(8)
        else:
            services = table.cell(3, 1).paragraphs[0].add_run(self.sp_visa_entry.toPlainText())
            services.font.name = ' In black '
            services.font.size = Pt(8)

        name_cust = table.cell(5, 1).paragraphs[0].add_run(self.name_edt.toPlainText())
        name_cust.font.name = ' In black '
        name_cust.font.size = Pt(8)

        if self.nat_cbx.currentText() != 'Other':
            nat = table.cell(6, 1).paragraphs[0].add_run(self.nat_cbx.currentText())
            nat.font.name = ' In black '
            nat.font.size = Pt(8)
        else:
            nat = table.cell(6, 1).paragraphs[0].add_run(self.ctry_entry.toPlainText())
            nat.font.name = ' In black '
            nat.font.size = Pt(8)

        fees = table.cell(8, 1).paragraphs[0].add_run(fee)
        fees.font.name = ' In black '
        fees.font.size = Pt(8)

        surcharge = table.cell(9, 1).paragraphs[0].add_run(icwf)
        surcharge.font.name = ' In black '
        surcharge.font.size = Pt(8)

        misc_amt = table.cell(10, 1).paragraphs[0].add_run(misc)
        misc_amt.font.name = ' In black '
        misc_amt.font.size = Pt(8)

        if self.misc_entry.value() != 0:
            misc_desc = table.cell(11, 1).paragraphs[0].add_run(oth_desc)
            misc_desc.font.name = ' In black '
            misc_desc.font.size = Pt(8)

        pmt_mode = table.cell(12, 1).paragraphs[0].add_run(self.trans_grp.checkedButton().text())
        pmt_mode.font.name = ' In black '
        pmt_mode.font.size = Pt(8)

        if self.postal_entry.value() != 0:
            post_chrg_lbl = table.cell(13, 0).paragraphs[0].add_run('Postal Charges:')
            post_chrg_lbl.font.name = ' In black '
            post_chrg_lbl.font.size = Pt(8)
            post_amt = table.cell(13, 1).paragraphs[0].add_run(post)
            post_amt.font.name = ' In black '
            post_amt.font.size = Pt(8)

        doc_no = table.cell(14, 1).paragraphs[0].add_run(str(self.doc_entry.value()))
        doc_no.font.name = ' In black '
        doc_no.font.size = Pt(8)

        tot_amt = table.cell(16, 1).paragraphs[0].add_run(total)
        tot_amt.font.name = ' In black '
        tot_amt.font.size = Pt(8)

        docx.save(in_file)

        wdFormatPDF = 17

        comtypes.CoInitialize()
        word = comtypes.client.CreateObject('Word.Application', dynamic=True)
        word.Documents.Open(in_file)
        word.Documents[0].SaveAs(out_file, wdFormatPDF)
        word.Documents[0].Close()
        word.Quit()

        if os.path.exists(in_file):
            os.remove(in_file)

        returncode = webbrowser.open_new(out_file + ".pdf")
        if returncode:
            self.clear_Ui()

            self.rcpt_entry.setText(str(int(receiptNo) + 1))

            self.prnt_btn.setEnabled(False)
            self.edt_btn.setEnabled(False)
            self.edt_btn.setVisible(False)
            self.save_btn.setVisible(False)

    # Disable all edit fields
    def save_Ui(self):
        self.name_edt.setEnabled(False)
        self.cash_rbtn.setEnabled(False)
        self.bank_rbtn.setEnabled(False)
        self.card_rbtn.setEnabled(False)
        self.gratis_rbtn.setEnabled(False)
        self.Rwrkday_rbtn.setEnabled(False)
        self.Rholiday_rbtn.setEnabled(False)
        self.pExpress_chk.setEnabled(False)
        self.wave_chk.setEnabled(False)
        self.other_entry.setEnabled(False)
        self.misc_entry.setEnabled(False)
        self.postal_entry.setEnabled(False)
        self.fees_entry.setEnabled(False)
        self.icwf_entry.setEnabled(False)
        self.prnt_btn.setEnabled(True)
        self.clear_btn.setEnabled(False)
        self.sp_visa_entry.setEnabled(False)
        self.ctry_entry.setEnabled(False)

    # Enable all fields
    def edit_Ui(self):
        self.srvt_cbx.setEnabled(True)
        self.edt_btn.setVisible(False)
        self.save_btn.setVisible(True)
        self.save_btn.setEnabled(True)
        self.clear_btn.setEnabled(True)
        self.prnt_btn.setEnabled(False)
        self.nat_cbx.setEnabled(True)
        self.name_edt.setEnabled(True)
        self.doc_entry.setEnabled(True)
        self.cash_rbtn.setEnabled(True)
        self.bank_rbtn.setEnabled(True)
        self.card_rbtn.setEnabled(True)
        self.gratis_rbtn.setEnabled(True)
        self.Rwrkday_rbtn.setEnabled(True)
        self.Rholiday_rbtn.setEnabled(True)
        self.pExpress_chk.setEnabled(True)
        self.wave_chk.setEnabled(True)
        self.other_entry.setEnabled(True)
        self.misc_entry.setEnabled(True)
        self.postal_entry.setEnabled(True)
        self.fees_entry.setEnabled(True)
        self.icwf_entry.setEnabled(True)

        if self.sp_visa_entry.toPlainText() != '':
            self.sp_visa_entry.setEnabled(True)

        if self.ctry_entry.toPlainText() != '':
            self.ctry_entry.setEnabled(True)

        if self.srvt_cbx.currentText() != "Choose Type":
            self.srvc_cbx.setEnabled(True)

    # Clear all data from fields
    def clear_Ui(self):
        self.name_edt.setText("")
        self.doc_entry.setValue(1)
        self.srvc_cbx.clear()
        self.srvc_cbx.setEnabled(False)
        self.srvt_cbx.setCurrentIndex(0)
        self.nat_cbx.setCurrentIndex(0)
        self.srvc_entry.setText("")
        self.ctry_entry.setText("")
        self.sp_visa_entry.setText("")
        if self.gratis_rbtn.isChecked():
            self.gratis_rbtn.setChecked(False)
        self.reset_Ui()

    # Set all UI elements to it's default state
    def reset_Ui(self):
        self.other_entry.setText("")
        self.misc_entry.setValue(0)
        self.postal_entry.setValue(0)
        self.fees_entry.setValue(0)
        self.icwf_entry.setValue(0)
        self.total_entry.setValue(0)
        self.pExpress_chk.setChecked(False)
        self.wave_chk.setChecked(False)

        self.bank_rbtn.setChecked(True)

        self.set_RWday()
        self.set_RHday()


    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Print Receipt"))
        self.prnt_btn.setText(_translate("MainWindow", "Print"))
        self.save_btn.setText(_translate("MainWindow", "Save"))
        self.clear_btn.setText(_translate("MainWindow", "Clear"))
        self.receipt_lbl.setText(_translate("MainWindow", "Receipt No:"))
        self.trns_lbl.setText(_translate("MainWindow", "Transaction Type:"))
        self.srvc_chrg_lbl.setText(_translate("MainWindow", "Service Charges:"))
        self.name_lbl.setText(_translate("MainWindow", "Name:"))
        self.srvc_lbl.setText(_translate("MainWindow", "Service:"))
        self.doc_lbl.setText(_translate("MainWindow", "Documents:"))
        self.label_3.setText(_translate("MainWindow", "Nationality:"))
        self.nat_cbx.setItemText(0, _translate("MainWindow", "Select Country"))
        self.total_lbl.setText(_translate("MainWindow", "Total"))
        self.post_lbl.setText(_translate("MainWindow", "Postal cost"))
        self.misc_lbl.setText(_translate("MainWindow", "Misc."))
        self.icwf_lbl.setText(_translate("MainWindow", "ICWF"))
        self.fees_lbl.setText(_translate("MainWindow", "Fees"))
        self.other_lbl.setText(_translate("MainWindow", "Misc Description:"))
        self.date_lbl.setText(_translate("MainWindow", "Date:"))
        self.label_10.setText(
            _translate("MainWindow", "<html><head/><body><p>"
                                     "<img src=\"./assets/Munich_logo.png\"/>"
                                     "</p></body></html>"))
        self.pExpress_chk.setText(_translate("MainWindow", "Postal express"))
        self.wave_chk.setText(_translate("MainWindow", "Wave ICWF"))
        self.Rwrkday_rbtn.setText(_translate("MainWindow", "R Work Day"))
        self.gratis_rbtn.setText(_translate("MainWindow", "Gratis"))
        self.Rholiday_rbtn.setText(_translate("MainWindow", "R Holiday"))
        self.edt_btn.setText(_translate("MainWindow", "Edit"))
        self.srvt_cbx.setItemText(0, _translate("MainWindow", "Service Type"))
        self.card_rbtn.setText(_translate("MainWindow", "Card"))
        self.cash_rbtn.setText(_translate("MainWindow", "Cash"))
        self.bank_rbtn.setText(_translate("MainWindow", "Bank Transfer"))
        self.email_id.setText(_translate("MainWindow", "Website: cgimunich.gov.in"))
        self.exit_btn.setText(_translate("MainWindow", "Exit"))
        self.plus.setText(_translate("MainWindow", "+"))
        self.plus_2.setText(_translate("MainWindow", "+"))
        self.plus_3.setText(_translate("MainWindow", "+"))
        self.equal.setText(_translate("MainWindow", "="))


if __name__ == "__main__":
    import sys

    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
