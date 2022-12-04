# Developer: Shaik Riyaz
# Created for: CGI Munich

import os
import re
import time
import webbrowser
import pandas as pd
import comtypes.client
from datetime import date
from docx import Document
from collections import OrderedDict
from PyQt5 import QtCore, QtGui, QtWidgets

import db_utils as db_conn
import Message_DialogBox as DialogBox


class Ui_ReportWindow(object):

    def __init__(self):
        self.exg_rate = 0

    def setupUi(self, ReportWindow):
        ReportWindow.setObjectName("ReportWindow")
        ReportWindow.setEnabled(True)
        ReportWindow.resize(1020, 800)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(ReportWindow.sizePolicy().hasHeightForWidth())
        ReportWindow.setSizePolicy(sizePolicy)
        ReportWindow.setMinimumSize(QtCore.QSize(1020, 760))
        ReportWindow.setMaximumSize(QtCore.QSize(1020, 760))
        self.centralwidget = QtWidgets.QWidget(ReportWindow)
        self.centralwidget.setEnabled(True)
        self.centralwidget.setObjectName("centralwidget")
        self.print_dialog = QtWidgets.QDialog(self.centralwidget)
        self.print_dialog.setObjectName("print_dialog")
        self.print_dialog.resize(472, 168)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.print_dialog.sizePolicy().hasHeightForWidth())
        self.print_dialog.setSizePolicy(sizePolicy)
        self.print_dialog.setMinimumSize(QtCore.QSize(472, 168))
        self.print_dialog.setMaximumSize(QtCore.QSize(472, 168))
        self.header_tags = ['Payment Receipt No.', 'Payment Date', 'Payment Time', 'Refund Receipt No.', 'Refund Date',
                            'Refund Time', 'Full Name', 'Nationality', 'No Docs', 'Service Code', 'Service Type',
                            'Service Sub Type', 'Service Description', 'Transaction type', 'Special Service', 'Gratis',
                            'Postal Express', 'Wave ICWF', 'Misc Desc', 'Misc Amount', 'Postal Charges', 'Fees',
                            'ICWF Charges', 'Total Amount', 'Total Refund Amount']
        self.line_16 = QtWidgets.QFrame(self.centralwidget)
        self.line_16.setGeometry(QtCore.QRect(10, 0, 1001, 20))
        self.line_16.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_16.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_16.setObjectName("line_16")
        self.line_2 = QtWidgets.QFrame(self.centralwidget)
        self.line_2.setGeometry(QtCore.QRect(0, 10, 20, 721))
        self.line_2.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_2.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_2.setObjectName("line_2")
        self.line_3 = QtWidgets.QFrame(self.centralwidget)
        self.line_3.setGeometry(QtCore.QRect(1000, 10, 20, 721))
        self.line_3.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_3.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_3.setObjectName("line_3")
        self.line_17 = QtWidgets.QFrame(self.centralwidget)
        self.line_17.setGeometry(QtCore.QRect(10, 690, 1001, 20))
        self.line_17.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_17.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_17.setObjectName("line_17")
        self.label_10 = QtWidgets.QLabel(self.centralwidget)
        self.label_10.setGeometry(QtCore.QRect(280, 10, 501, 71))
        self.label_10.setAlignment(QtCore.Qt.AlignCenter)
        self.label_10.setObjectName("label_10")
        self.email_id = QtWidgets.QLabel(self.centralwidget)
        self.email_id.setGeometry(QtCore.QRect(370, 80, 211, 21))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(11)
        self.email_id.setFont(font)
        self.email_id.setStyleSheet("color: rgb(81, 81, 81);")
        self.email_id.setAlignment(QtCore.Qt.AlignLeading | QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter)
        self.email_id.setObjectName("email_id")
        self.line_18 = QtWidgets.QFrame(self.centralwidget)
        self.line_18.setGeometry(QtCore.QRect(10, 100, 1001, 20))
        self.line_18.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_18.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_18.setObjectName("line_18")
        self.horizontalLayoutWidget = QtWidgets.QWidget(self.centralwidget)
        self.horizontalLayoutWidget.setGeometry(QtCore.QRect(20, 20, 186, 80))
        self.horizontalLayoutWidget.setObjectName("horizontalLayoutWidget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget)
        self.horizontalLayout.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.date_lbl = QtWidgets.QLabel(self.horizontalLayoutWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(1)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.date_lbl.sizePolicy().hasHeightForWidth())
        self.date_lbl.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        self.date_lbl.setFont(font)
        self.date_lbl.setAlignment(QtCore.Qt.AlignCenter)
        self.date_lbl.setObjectName("date_lbl")
        self.horizontalLayout.addWidget(self.date_lbl)
        self.date_entry = QtWidgets.QLabel(self.horizontalLayoutWidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(2)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.date_entry.sizePolicy().hasHeightForWidth())
        self.date_entry.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setPointSize(14)
        self.date_entry.setFont(font)
        self.date_entry.setText("")
        self.date_entry.setAlignment(QtCore.Qt.AlignLeading | QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter)
        self.date_entry.setObjectName("date_entry")
        self.horizontalLayout.addWidget(self.date_entry)
        self.from_date = QtWidgets.QDateEdit(self.centralwidget)
        self.from_date.setGeometry(QtCore.QRect(100, 140, 211, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.from_date.setFont(font)
        self.from_date.setAlignment(QtCore.Qt.AlignCenter)
        self.from_date.setCalendarPopup(True)
        self.from_date.setObjectName("from_date")
        self.from_date.setDate(date.today())
        d = QtCore.QDate(date.today().year, 12, 31)
        self.from_date.setMaximumDate(d)
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(20, 150, 71, 21))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.label.setFont(font)
        self.label.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignTrailing | QtCore.Qt.AlignVCenter)
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(390, 150, 51, 20))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.label_2.setFont(font)
        self.label_2.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignTrailing | QtCore.Qt.AlignVCenter)
        self.label_2.setObjectName("label_2")
        self.to_date = QtWidgets.QDateEdit(self.centralwidget)
        self.to_date.setGeometry(QtCore.QRect(450, 140, 211, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.to_date.setFont(font)
        self.to_date.setAlignment(QtCore.Qt.AlignCenter)
        self.to_date.setCalendarPopup(True)
        self.to_date.setObjectName("to_date")
        self.to_date.setDate(date.today())
        self.print_btn = QtWidgets.QPushButton(self.centralwidget)
        self.print_btn.setGeometry(QtCore.QRect(20, 650, 150, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(False)
        font.setWeight(50)
        self.print_btn.setFont(font)
        self.print_btn.setObjectName("print_btn")
        self.line = QtWidgets.QFrame(self.centralwidget)
        self.line.setGeometry(QtCore.QRect(10, 200, 1001, 20))
        self.line.setFrameShape(QtWidgets.QFrame.HLine)
        self.line.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line.setObjectName("line")
        self.exit_btn = QtWidgets.QPushButton(self.centralwidget, clicked=lambda: ReportWindow.close())
        self.exit_btn.setGeometry(QtCore.QRect(850, 650, 150, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(False)
        font.setWeight(50)
        self.exit_btn.setFont(font)
        self.exit_btn.setObjectName("exit_btn")
        self.daily_rpt_cbx = QtWidgets.QCheckBox(self.centralwidget)
        self.daily_rpt_cbx.setGeometry(QtCore.QRect(780, 140, 211, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.daily_rpt_cbx.setFont(font)
        self.daily_rpt_cbx.setObjectName("daily_rpt_cbx")
        self.filterwidget = QtWidgets.QWidget(self.centralwidget)
        self.filterwidget.setEnabled(True)
        self.filterwidget.setGeometry(QtCore.QRect(9, 209, 1001, 431))
        self.filterwidget.setObjectName("filterwidget")
        self.line_9 = QtWidgets.QFrame(self.filterwidget)
        self.line_9.setGeometry(QtCore.QRect(330, 170, 20, 141))
        self.line_9.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_9.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_9.setObjectName("line_9")
        self.line_13 = QtWidgets.QFrame(self.filterwidget)
        self.line_13.setGeometry(QtCore.QRect(630, 180, 20, 231))
        self.line_13.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_13.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_13.setObjectName("line_13")
        self.line_8 = QtWidgets.QFrame(self.filterwidget)
        self.line_8.setGeometry(QtCore.QRect(0, 180, 20, 131))
        self.line_8.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_8.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_8.setObjectName("line_8")
        self.verticalLayoutWidget_2 = QtWidgets.QWidget(self.filterwidget)
        self.verticalLayoutWidget_2.setGeometry(QtCore.QRect(660, 180, 311, 221))
        self.verticalLayoutWidget_2.setObjectName("verticalLayoutWidget_2")
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self.verticalLayoutWidget_2)
        self.verticalLayout_2.setContentsMargins(0, 0, 0, 0)
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.gratis_sr_cbx = QtWidgets.QCheckBox(self.verticalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.gratis_sr_cbx.setFont(font)
        self.gratis_sr_cbx.setObjectName("gratis_sr_cbx")
        self.verticalLayout_2.addWidget(self.gratis_sr_cbx)
        self.rworkday_sr_cbx = QtWidgets.QCheckBox(self.verticalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.rworkday_sr_cbx.setFont(font)
        self.rworkday_sr_cbx.setObjectName("rworkday_sr_cbx")
        self.verticalLayout_2.addWidget(self.rworkday_sr_cbx)
        self.rholiday_sr_cbx = QtWidgets.QCheckBox(self.verticalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.rholiday_sr_cbx.setFont(font)
        self.rholiday_sr_cbx.setObjectName("rholiday_sr_cbx")
        self.verticalLayout_2.addWidget(self.rholiday_sr_cbx)
        self.postal_sr_cbx = QtWidgets.QCheckBox(self.verticalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.postal_sr_cbx.setFont(font)
        self.postal_sr_cbx.setObjectName("postal_sr_cbx")
        self.verticalLayout_2.addWidget(self.postal_sr_cbx)
        self.misc_sr_cbx = QtWidgets.QCheckBox(self.verticalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.misc_sr_cbx.setFont(font)
        self.misc_sr_cbx.setObjectName("misc_sr_cbx")
        self.verticalLayout_2.addWidget(self.misc_sr_cbx)
        self.refund_ind_cbx = QtWidgets.QCheckBox(self.verticalLayoutWidget_2)
        self.refund_ind_cbx.setObjectName(u"refund_ind_cbx")
        self.refund_ind_cbx.setFont(font)
        self.verticalLayout_2.addWidget(self.refund_ind_cbx)
        self.line_12 = QtWidgets.QFrame(self.filterwidget)
        self.line_12.setGeometry(QtCore.QRect(980, 170, 20, 241))
        self.line_12.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_12.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_12.setObjectName("line_12")
        self.verticalLayoutWidget = QtWidgets.QWidget(self.filterwidget)
        self.verticalLayoutWidget.setGeometry(QtCore.QRect(30, 180, 291, 121))
        self.verticalLayoutWidget.setObjectName("verticalLayoutWidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.verticalLayoutWidget)
        self.verticalLayout.setContentsMargins(0, 0, 0, 0)
        self.verticalLayout.setObjectName("verticalLayout")
        self.bank_tr_cbx = QtWidgets.QCheckBox(self.verticalLayoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.bank_tr_cbx.setFont(font)
        self.bank_tr_cbx.setObjectName("bank_tr_cbx")
        self.verticalLayout.addWidget(self.bank_tr_cbx)
        self.cash_tr_cbx = QtWidgets.QCheckBox(self.verticalLayoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.cash_tr_cbx.setFont(font)
        self.cash_tr_cbx.setObjectName("cash_tr_cbx")
        self.verticalLayout.addWidget(self.cash_tr_cbx)
        self.card_tr_cbx = QtWidgets.QCheckBox(self.verticalLayoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.card_tr_cbx.setFont(font)
        self.card_tr_cbx.setObjectName("card_tr_cbx")
        self.verticalLayout.addWidget(self.card_tr_cbx)
        self.line_11 = QtWidgets.QFrame(self.filterwidget)
        self.line_11.setGeometry(QtCore.QRect(220, 160, 121, 20))
        self.line_11.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_11.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_11.setObjectName("line_11")
        self.service_code_cbx = QtWidgets.QComboBox(self.filterwidget)
        self.service_code_cbx.setEnabled(True)
        self.service_code_cbx.setGeometry(QtCore.QRect(390, 80, 211, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.service_code_cbx.setFont(font)
        self.service_code_cbx.setEditable(True)
        self.service_code_cbx.setObjectName("service_code_cbx")
        self.line_10 = QtWidgets.QFrame(self.filterwidget)
        self.line_10.setGeometry(QtCore.QRect(10, 300, 331, 20))
        self.line_10.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_10.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_10.setObjectName("line_10")
        self.country_cbx = QtWidgets.QComboBox(self.filterwidget)
        self.country_cbx.setGeometry(QtCore.QRect(170, 20, 211, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.country_cbx.setFont(font)
        self.country_cbx.setEditable(True)
        self.country_cbx.setObjectName("country_cbx")
        self.country_cbx.addItem("")
        self.country_cbx.addItem("")
        self.trns_lbl = QtWidgets.QLabel(self.filterwidget)
        self.trns_lbl.setGeometry(QtCore.QRect(10, 150, 211, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        self.trns_lbl.setFont(font)
        self.trns_lbl.setObjectName("trns_lbl")
        self.line_14 = QtWidgets.QFrame(self.filterwidget)
        self.line_14.setGeometry(QtCore.QRect(640, 400, 351, 20))
        self.line_14.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_14.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_14.setObjectName("line_14")
        self.srvc_chrg_lbl = QtWidgets.QLabel(self.filterwidget)
        self.srvc_chrg_lbl.setGeometry(QtCore.QRect(640, 150, 201, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        self.srvc_chrg_lbl.setFont(font)
        self.srvc_chrg_lbl.setObjectName("srvc_chrg_lbl")
        self.service_subType_cbx = QtWidgets.QComboBox(self.filterwidget)
        self.service_subType_cbx.setEnabled(True)
        self.service_subType_cbx.setGeometry(QtCore.QRect(610, 80, 311, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.service_subType_cbx.setFont(font)
        self.service_subType_cbx.setEditable(True)
        self.service_subType_cbx.setObjectName("service_subType_cbx")
        self.country_lbl = QtWidgets.QLabel(self.filterwidget)
        self.country_lbl.setGeometry(QtCore.QRect(30, 20, 131, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.country_lbl.setFont(font)
        self.country_lbl.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignTrailing | QtCore.Qt.AlignVCenter)
        self.country_lbl.setObjectName("country_lbl")
        self.label_3 = QtWidgets.QLabel(self.filterwidget)
        self.label_3.setGeometry(QtCore.QRect(70, 80, 91, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.label_3.setFont(font)
        self.label_3.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignTrailing | QtCore.Qt.AlignVCenter)
        self.label_3.setObjectName("label_3")
        self.line_15 = QtWidgets.QFrame(self.filterwidget)
        self.line_15.setGeometry(QtCore.QRect(840, 160, 151, 20))
        self.line_15.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_15.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_15.setObjectName("line_15")
        self.service_type_cbx = QtWidgets.QComboBox(self.filterwidget)
        self.service_type_cbx.setGeometry(QtCore.QRect(170, 80, 211, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.service_type_cbx.setFont(font)
        self.service_type_cbx.setEditable(True)
        self.service_type_cbx.setObjectName("service_type_cbx")
        self.service_type_cbx.addItem("")
        self.line_19 = QtWidgets.QFrame(self.centralwidget)
        self.line_19.setGeometry(QtCore.QRect(10, 720, 1001, 20))
        self.line_19.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_19.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_19.setObjectName("line_19")

        self.label_4 = QtWidgets.QLabel(self.print_dialog)
        self.label_4.setGeometry(QtCore.QRect(10, 0, 161, 51))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(True)
        font.setWeight(75)
        self.label_4.setFont(font)
        self.label_4.setObjectName("label_4")
        self.line_21 = QtWidgets.QFrame(self.print_dialog)
        self.line_21.setGeometry(QtCore.QRect(10, 100, 451, 20))
        self.line_21.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_21.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_21.setObjectName("line_21")
        self.line_19 = QtWidgets.QFrame(self.print_dialog)
        self.line_19.setGeometry(QtCore.QRect(90, 20, 371, 20))
        self.line_19.setFrameShape(QtWidgets.QFrame.HLine)
        self.line_19.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_19.setObjectName("line_19")
        self.line_20 = QtWidgets.QFrame(self.print_dialog)
        self.line_20.setGeometry(QtCore.QRect(0, 40, 20, 71))
        self.line_20.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_20.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_20.setObjectName("line_20")
        self.horizontalLayoutWidget_2 = QtWidgets.QWidget(self.print_dialog)
        self.horizontalLayoutWidget_2.setGeometry(QtCore.QRect(30, 40, 411, 61))
        self.horizontalLayoutWidget_2.setObjectName("horizontalLayoutWidget_2")
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget_2)
        self.horizontalLayout_2.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        self.print_pdf_cbx = QtWidgets.QCheckBox(self.horizontalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.print_pdf_cbx.setFont(font)
        self.print_pdf_cbx.setChecked(True)
        self.print_pdf_cbx.setObjectName("print_pdf_cbx")
        self.horizontalLayout_2.addWidget(self.print_pdf_cbx)
        self.print_csv_cbx = QtWidgets.QCheckBox(self.horizontalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.print_csv_cbx.setFont(font)
        self.print_csv_cbx.setObjectName("print_csv_cbx")
        self.horizontalLayout_2.addWidget(self.print_csv_cbx)
        self.print_excel_cbx = QtWidgets.QCheckBox(self.horizontalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.print_excel_cbx.setFont(font)
        self.print_excel_cbx.setObjectName("print_excel_cbx")
        self.horizontalLayout_2.addWidget(self.print_excel_cbx)
        self.line_22 = QtWidgets.QFrame(self.print_dialog)
        self.line_22.setGeometry(QtCore.QRect(450, 30, 20, 81))
        self.line_22.setFrameShape(QtWidgets.QFrame.VLine)
        self.line_22.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line_22.setObjectName("line_22")
        self.print_open_btn = QtWidgets.QPushButton(self.print_dialog)
        self.print_open_btn.setGeometry(QtCore.QRect(10, 120, 150, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(False)
        font.setWeight(50)
        self.print_open_btn.setFont(font)
        self.print_open_btn.setStyleSheet("")
        self.print_open_btn.setObjectName("print_open_btn")
        self.cancel_print_btn = QtWidgets.QPushButton(self.print_dialog)
        self.cancel_print_btn.setGeometry(QtCore.QRect(310, 120, 150, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(False)
        font.setWeight(50)
        self.cancel_print_btn.setFont(font)
        self.cancel_print_btn.setStyleSheet("")
        self.cancel_print_btn.setObjectName("cancel_print_btn")
        ReportWindow.setCentralWidget(self.centralwidget)
        self.statusbar = QtWidgets.QStatusBar(ReportWindow)
        self.statusbar.setObjectName("statusbar")
        ReportWindow.setStatusBar(self.statusbar)

        self.retranslateUi(ReportWindow)
        QtCore.QMetaObject.connectSlotsByName(ReportWindow)

        # Function calls
        self.date_assign()
        self.get_country_list()
        self.get_service_types()
        self.service_type_cbx.currentIndexChanged.connect(self.get_service_codes_subtypes)
        self.get_service_codes_subtypes(self.service_type_cbx.currentIndex())

        self.daily_rpt_cbx.stateChanged.connect(lambda: self.btnstate(self.daily_rpt_cbx))

        self.print_btn.clicked.connect(self.print_report)

        self.print_open_btn.clicked.connect(self.get_saveOpts)
        self.cancel_print_btn.clicked.connect(self.print_dialog.close)

    # Date assign
    def date_assign(self):
        self.date_entry.setText(str(date.today()))

    # Get Nationality
    def get_country_list(self):

        selectquery = "SELECT name FROM nationality"
        records = db_conn.fetch_rows(selectquery, None)
        if records:
            for row in records:
                self.country_cbx.addItem(row[0])
        else:
            DialogBox.errordialog('Nationality database is empty.')

    # Get Service Type
    def get_service_types(self):
        items = []
        items1 = []
        selectquery = "SELECT code, category, sub_category FROM services"
        records = db_conn.fetch_rows(selectquery, None)
        if records:
            for row in records:
                items.append(str(row[0]) + ":" + row[2])
                items1.append(row[1])
            k = sorted(set(items1))
            serv_dict = dict(zip(items, items1))
            for i in k:
                final_key = ["Select Code:Select SubType"]
                for key in serv_dict:
                    if i == serv_dict[key]:
                        final_key.append(key)
                self.service_type_cbx.addItem(i, final_key)

        else:
            DialogBox.errordialog('Services database is empty')

    # Generate service codes based on service type
    def get_service_codes_subtypes(self, index):
        if self.service_type_cbx.currentText() != "Select Type":
            self.service_code_cbx.setEnabled(True)
            self.service_code_cbx.clear()
            self.service_subType_cbx.setEnabled(True)
            self.service_subType_cbx.clear()
            code = self.service_type_cbx.itemData(index)
            codes = []
            sub_type = []
            for values in code:
                codes.append(values.split(":")[0])
                sub_type.append(values.split(":")[1])
                sub_type = list(OrderedDict.fromkeys(sub_type))
            if codes:
                self.service_code_cbx.addItems(codes)
                self.service_subType_cbx.addItems(sub_type)
        else:
            self.service_code_cbx.clear()
            self.service_code_cbx.setEnabled(False)
            self.service_subType_cbx.clear()
            self.service_subType_cbx.setEnabled(False)

    # Generate service codes based on service type
    def get_service_subType(self, index):
        if self.service_type_cbx.currentText() != "Service Type":
            self.service_subType_cbx.setEnabled(True)
            self.service_subType_cbx.clear()
            sub_type = self.service_type_cbx.itemData(index).split(":")[1]
            if sub_type:
                self.service_subType_cbx.addItems(sub_type)
        else:
            self.service_subType_cbx.clear()
            self.service_subType_cbx.setEnabled(False)

    def get_trans_type_query(self):
        where_clause = ""
        values = []
        if self.bank_tr_cbx.isChecked():
            where_clause += "receipt.trans_type LIKE %s OR "
            values.append(self.bank_tr_cbx.text())
        if self.cash_tr_cbx.isChecked():
            where_clause += "receipt.trans_type LIKE %s OR "
            values.append(self.cash_tr_cbx.text())
        if self.card_tr_cbx.isChecked():
            where_clause += "receipt.trans_type LIKE %s OR "
            values.append(self.card_tr_cbx.text())
        return where_clause, values

    def get_service_charges_query(self):
        where_clause = ""
        values = []
        if self.gratis_sr_cbx.isChecked():
            where_clause += "receipt.gratis_ind LIKE %s AND "
            values.append(self.gratis_sr_cbx.text())
        if self.rworkday_sr_cbx.isChecked():
            where_clause += "receipt.serv_chrg_type LIKE %s OR "
            values.append(self.rworkday_sr_cbx.text())
        if self.rholiday_sr_cbx.isChecked():
            where_clause += "receipt.serv_chrg_type LIKE %s OR "
            values.append(self.rholiday_sr_cbx.text())
        return where_clause, values

    def get_other_query(self):
        where_clause = ""
        values = []
        if self.postal_sr_cbx.isChecked():
            where_clause += "post_amt NOT LIKE %s AND "
            values.append("0")
        if self.misc_sr_cbx.isChecked():
            where_clause += "misc_amt NOT LIKE %s AND "
            values.append("0")
        if self.refund_ind_cbx.isChecked():
            where_clause += "refund_receipt_no NOT LIKE %s AND "
            values.append("NULL")
        return where_clause, values

    def create_whereclause(self):
        where_clause = ""
        values = []
        if self.country_cbx.currentText() != "Select Country":
            where_clause += "receipt.nationality LIKE %s AND "
            values.append(self.country_cbx.currentText())
        if self.service_type_cbx.currentText() != "Select Type":
            where_clause += "services.category LIKE %s AND "
            values.append(self.service_type_cbx.currentText())
            if self.service_code_cbx.currentText() != "Select Code":
                where_clause += "services.code LIKE %s AND"
                values.append(self.service_code_cbx.currentText())
            if self.service_subType_cbx.currentText() != "Select SubType":
                where_clause += "services.sub_category LIKE %s AND "
                values.append(self.service_subType_cbx.currentText())
        return where_clause, values

    def daily_report(self, records, refund_records):
        flag = True
        temp_dir = os.getcwd() + r'\templates\\'
        out_dir = str(os.path.expanduser('~')) + r'\Documents\reports\\'

        # Create reports folder
        if not os.path.exists(out_dir):
            os.makedirs(out_dir)

        in_file = out_dir + "daily_report.docx"

        if self.from_date.text() != self.to_date.text():
            out_file = out_dir + "DailyReport_" + self.from_date.text() + "_to_" + self.to_date.text()
        else:
            out_file = out_dir + "DailyReport_" + self.from_date.text()

        docx = Document(temp_dir + 'Daily_Account_Report.docx')

        total_eur, total_ref_eur, total_ppt, total_visa, total_oci, total_att, total_reg, total_crt, total_bnk, \
        total_misc, total_post, total_cash, total_card = 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0

        visa_bnk, visa_card, visa_cash, ppt_bnk, ppt_card, ppt_cash, oci_bnk, oci_card, oci_cash, \
        reg_bnk, reg_cash, reg_card, att_bnk, att_cash, att_card, misc_bnk, misc_cash, misc_card, \
        crt_bnk, crt_cash, crt_card, post_bnk, post_cash, post_card = 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, \
                                                                      0, 0, 0, 0, 0, 0, 0, 0

        total_ppt_grts, total_visa_grts, total_oci_grts, total_att_grts, total_reg_grts, total_crt_grts = 0, 0, 0, 0, 0, 0

        receipt_nos = []
        refund_receipt_nos = []

        if records:
            for row in records:
                total_eur += row[17]
                if row[5].find("PPT") != -1:
                    total_ppt += row[15] + row[16]
                    if row[7].find("Bank") != -1:
                        ppt_bnk += row[15] + row[16]
                    elif row[7].find("Cash") != -1:
                        ppt_cash += row[15] + row[16]
                    elif row[7].find("Card") != -1:
                        ppt_card += row[15] + row[16]

                    if row[9].find("Non Gratis") == -1:
                        total_ppt_grts += 1

                elif row[5].find("GVAL") != -1 or row[5].find("SV00") != -1:
                    total_visa += row[15] + row[16]
                    if row[7].find("Bank") != -1:
                        visa_bnk += row[15] + row[16]
                    elif row[7].find("Cash") != -1:
                        visa_cash += row[15] + row[16]
                    elif row[7].find("Card") != -1:
                        visa_card += row[15] + row[16]

                    if row[9].find("Non Gratis") == -1:
                        total_visa_grts += 1

                elif row[5].find("OCI") != -1:
                    total_oci += row[15] + row[16]
                    if row[7].find("Bank") != -1:
                        oci_bnk += row[15] + row[16]
                    elif row[7].find("Cash") != -1:
                        oci_cash += row[15] + row[16]
                    elif row[7].find("Card") != -1:
                        oci_card += row[15] + row[16]

                    if row[9].find("Non Gratis") == -1:
                        total_oci_grts += 1

                elif row[5].find("ATT") != -1:
                    total_att += row[15] + row[16]
                    if row[7].find("Bank") != -1:
                        att_bnk += row[15] + row[16]
                    elif row[7].find("Cash") != -1:
                        att_cash += row[15] + row[16]
                    elif row[7].find("Card") != -1:
                        att_card += row[15] + row[16]

                    if row[9].find("Non Gratis") == -1:
                        total_att_grts += 1

                elif row[5].find("CRT") != -1:
                    total_crt += row[15] + row[16]
                    if row[7].find("Bank") != -1:
                        crt_bnk += row[15] + row[16]
                    elif row[7].find("Cash") != -1:
                        crt_cash += row[15] + row[16]
                    elif row[7].find("Card") != -1:
                        crt_card += row[15] + row[16]

                    if row[9].find("Non Gratis") == -1:
                        total_crt_grts += 1

                elif row[5].find("REG") != -1:
                    total_reg += row[15] + row[16]
                    if row[7].find("Bank") != -1:
                        reg_bnk += row[15] + row[16]
                    elif row[7].find("Cash") != -1:
                        reg_cash += row[15] + row[16]
                    elif row[7].find("Card") != -1:
                        reg_card += row[15] + row[16]

                    if row[9].find("Non Gratis") == -1:
                        total_reg_grts += 1

                if row[12] != 0:
                    total_misc += row[12]
                    if row[7].find("Bank") != -1:
                        misc_bnk += row[12]
                    elif row[7].find("Cash") != -1:
                        misc_cash += row[12]
                    elif row[7].find("Card") != -1:
                        misc_card += row[12]

                if row[14] != 0:
                    total_post += row[14]
                    if row[7].find("Bank") != -1:
                        post_bnk += row[14]
                    elif row[7].find("Cash") != -1:
                        post_cash += row[14]
                    elif row[7].find("Card") != -1:
                        post_card += row[14]

                if row[7].find("Bank") != -1:
                    total_bnk += row[17]
                elif row[7].find("Cash") != -1:
                    total_cash += row[17]
                elif row[7].find("Card") != -1:
                    total_card += row[17]

                receipt_nos.append(row[0])
        else:
            receipt_nos.append("--")

        if refund_records:
            for row in refund_records:
                refund_receipt_nos.append(row[0])
                total_ref_eur += row[1]
        else:
            refund_receipt_nos.append("--")

        total_inr = round(total_eur * self.exg_rate, 2)
        total_ref_inr = round(total_ref_eur * self.exg_rate, 2)

        table_values = [total_ppt, total_visa, total_oci, total_att, total_reg, total_crt, total_misc, total_post,
                        total_cash, round(total_cash * self.exg_rate, 2), total_bnk,
                        round(total_bnk * self.exg_rate, 2),
                        total_card, round(total_card * self.exg_rate, 2), ppt_cash, ppt_bnk, ppt_card, visa_cash,
                        visa_bnk, visa_card, oci_cash, oci_bnk, oci_card, att_cash, att_bnk, att_card, reg_cash,
                        reg_bnk, reg_card, crt_cash, crt_bnk, crt_card, misc_cash, misc_bnk, misc_card, post_cash,
                        post_bnk, post_card,
                        "Gratis/" + str(total_ppt_grts) if total_ppt_grts != 0 else " ",
                        "Gratis/" + str(total_visa_grts) if total_visa_grts != 0 else " ",
                        "Gratis/" + str(total_oci_grts) if total_oci_grts != 0 else " ",
                        "Gratis/" + str(total_att_grts) if total_att_grts != 0 else " ",
                        "Gratis/" + str(total_reg_grts) if total_reg_grts != 0 else " ",
                        "Gratis/" + str(total_crt_grts) if total_crt_grts != 0 else " ",
                        total_ref_eur]

        regex = re.compile("total_eur")
        regex1 = re.compile("total_inr")
        regex2 = re.compile("total_ref_eur")
        regex3 = re.compile("total_ref_inr")

        regex4 = re.compile("fdate")
        regex5 = re.compile("tdate")
        regex6 = re.compile("rno")
        regex7 = re.compile("r1no")
        regex8 = re.compile("r2no")
        regex9 = re.compile("r3no")

        regex10 = re.compile("dat_value")
        regex11 = re.compile("ex_value")

        self.docx_replace_regex(docx, regex, "€ " + str(total_eur))
        self.docx_replace_regex(docx, regex1, "₹ " + str(total_inr))
        self.docx_replace_regex(docx, regex2, "€ " + str(total_ref_eur))
        self.docx_replace_regex(docx, regex3, "₹ " + str(total_ref_inr))

        self.docx_replace_regex(docx, regex4, self.from_date.text())
        self.docx_replace_regex(docx, regex5, self.to_date.text())
        self.docx_replace_regex(docx, regex6, str(receipt_nos[0]))
        self.docx_replace_regex(docx, regex7, str(receipt_nos[len(receipt_nos) - 1]))
        self.docx_replace_regex(docx, regex8, str(refund_receipt_nos[0]))
        self.docx_replace_regex(docx, regex9, str(refund_receipt_nos[len(refund_receipt_nos) - 1]))

        self.docx_replace_regex(docx, regex10, self.date_entry.text())
        self.docx_replace_regex(docx, regex11, str(self.exg_rate))

        j = 1
        for data in table_values:
            regex12 = re.compile("Value" + str(j) + "\$")
            self.docx_replace_regex(docx, regex12, str(data))
            j += 1

        if os.path.exists(out_file + '.pdf'):
            try:
                os.remove(out_file + '.pdf')
            except IOError:
                flag = False
                DialogBox.errordialog("PDF file is in use. Please close the file and print again.")

        if flag:
            docx.save(in_file)
            wdFormatPDF = 17
            comtypes.CoInitialize()
            word = comtypes.client.CreateObject('Word.Application', dynamic=True)
            word.Documents.Open(in_file)
            word.Documents[0].SaveAs(out_file, wdFormatPDF)
            word.Documents[0].Close()
            word.Quit()
            if os.path.exists(in_file):
                os.remove(in_file)
            returncode = webbrowser.open_new(out_file + ".pdf")
            return returncode
        else:
            return flag

    def pdf_generate_report(self):
        pdf_returncode = False
        columns = ""

        refund_find_query = "SELECT receipt_no, total_refund FROM refund_receipt " \
                            "WHERE refund_receipt.receipt_date BETWEEN %s AND %s"
        value = [self.from_date.text(), self.to_date.text()]
        values = tuple(value, )
        ref_records = db_conn.fetch_rows(refund_find_query, values)

        find_query = "SELECT receipt.*, services.category, services.sub_category FROM receipt INNER JOIN services " \
                     "ON receipt.service_code LIKE services.code AND ( receipt.receipt_date BETWEEN %s AND %s )" + \
                     columns

        values = tuple(value, )
        records = db_conn.fetch_rows(find_query, values)

        if records or ref_records:
            pdf_returncode = self.daily_report(records, ref_records)
        else:
            DialogBox.errordialog('There is no data available from [' + self.from_date.text() + '] to ['
                                  + self.to_date.text() + ']')
        return pdf_returncode

    def csv_excel_generate_report(self, type, opt):
        flag = True
        csv_returncode = False
        excel_returncode = False
        columns = ""
        value = []

        out_dir = os.getcwd() + r'\files\\'
        if self.from_date.text() != self.to_date.text():
            csvout_fileName = out_dir + "Report_" + self.from_date.text() + "_to_" + self.to_date.text() + '.csv'
            excelout_fileName = out_dir + "Report_" + self.from_date.text() + "_to_" + self.to_date.text() + '.xlsx'
        else:
            csvout_fileName = out_dir + "Report_" + self.from_date.text() + '.csv'
            excelout_fileName = out_dir + "Report_" + self.from_date.text() + '.xlsx'

        if not type:
            columns = "AND (" + self.create_whereclause()[0][:-4] + ")" if self.create_whereclause()[0] != "" else ""
            columns += "AND (" + self.get_trans_type_query()[0][:-4] + ")" if self.get_trans_type_query()[
                                                                                  0] != "" else ""
            columns += "AND (" + self.get_other_query()[0][:-4] + ")" if self.get_other_query()[
                                                                             0] != "" else ""
            columns += "AND (" + self.get_service_charges_query()[0][:-4] + ")" if self.get_service_charges_query()[
                                                                                       0] != "" else ""

            value1 = self.create_whereclause()[1]
            value2 = self.get_trans_type_query()[1]
            value3 = self.get_other_query()[1]
            value4 = self.get_service_charges_query()[1]

            value = value1 + value2 + value3 + value4

        dates = [self.from_date.text(), self.to_date.text()]
        final_values = dates + dates + value

        find_query = "SELECT receipt.receipt_no, receipt.receipt_date, receipt.receipt_time, " \
                     "IFNULL(refund_receipt.receipt_no,'None'), " \
                     "IFNULL(refund_receipt.receipt_date,'None'), " \
                     "IFNULL(refund_receipt.receipt_time,'None'), " \
                     "receipt.Name, receipt.Nationality, receipt.no_doc, services.code, services.category, " \
                     "services.sub_category, services.description, receipt.trans_type, receipt.serv_chrg_type, " \
                     "receipt.gratis_ind, receipt.post_express, receipt.wave_icwf, receipt.misc_desc, receipt.misc_amt, " \
                     "receipt.post_amt, receipt.fees_amt, receipt.icwf_amt, receipt.total_amt, " \
                     "IFNULL(refund_receipt.total_refund,'') FROM receipt " \
                     "INNER JOIN services ON receipt.service_code LIKE services.code " \
                     "LEFT OUTER JOIN refund_receipt ON refund_receipt.receipt_no LIKE receipt.refund_receipt_no " \
                     "WHERE ((refund_receipt.receipt_date BETWEEN %s AND %s) OR (receipt.receipt_date BETWEEN %s AND %s)) " \
                     + columns
        values = tuple(final_values, )
        records = db_conn.fetch_rows(find_query, values)

        if records:
            df = pd.DataFrame(records)
            df.columns = self.header_tags

            paymentseconds = df['Payment Time'].dt.total_seconds()
            df['Payment Time'] = self.convert_time(paymentseconds).values

            if opt == "csv":
                try:
                    df.to_csv(csvout_fileName, index=False)
                    csv_returncode = True
                except IOError:
                    DialogBox.errordialog("csv file is in use. Please close the file and try again.")

            if opt == "excel":
                try:
                    df.to_excel(excelout_fileName, sheet_name='Sheet1', index=False)
                    excel_returncode = True
                except IOError:
                    DialogBox.errordialog("Excel file is in use. Please close the file and try again.")
        else:
            DialogBox.errordialog('There is no data available for [' + ','.join(value) + "]")

        return csv_returncode, excel_returncode

    def convert_time(self, seconds):
        mlist = []
        for i in seconds:
            mlist.append(time.strftime("%H:%M:%S", time.gmtime(i)))

        df = pd.DataFrame({'Time': mlist})

        return df['Time']

    def field_validations(self):
        msg = ""
        if self.country_cbx.currentText() == "":
            msg += "Nationality Field cannot be blank!!!\n"
        if self.service_type_cbx.currentText() == "":
            msg += "Service Type Field cannot be blank!!!\n"
        else:
            if self.service_code_cbx.isEnabled() and self.service_code_cbx.currentText() == "":
                msg += "Service Code Field cannot be blank!!!\n"
            elif self.service_subType_cbx.isEnabled() and self.service_subType_cbx.currentText() == "":
                msg += "Service SubType cannot be blank!!!\n"
        return msg

    def btnstate(self, daily_rpt_cbx):
        if daily_rpt_cbx.isChecked():
            self.filterwidget.setEnabled(False)
        else:
            self.filterwidget.setEnabled(True)

    def print_report(self):
        rpdf = False
        rcsv = False
        rexcel = False
        self.saveOpt = []
        errors = ""

        if self.from_date.date().month() != self.to_date.date().month():
            errors += "Invalid Dates. Report can only be generated within the same month."
        elif self.from_date.date() > self.to_date.date():
            errors += "Invalid Dates. From date cannot be greater than To date"
        elif not self.getExchangeRate():
            errors += 'There is no exchange rate available for the provided month and year'

        if not self.daily_rpt_cbx.isChecked() and errors == "":
            errors = self.field_validations()

        if errors == "":
            self.showSaveDialogBox()
            if self.saveOpt:
                for opts in self.saveOpt:
                    if opts == "pdf":
                        rpdf = self.pdf_generate_report()
                    if opts in ["csv", "excel"]:
                        rcsv, rexcel = self.csv_excel_generate_report(self.daily_rpt_cbx.isChecked(), opts)
                if rpdf or rcsv or rexcel:
                    DialogBox.infodialog("Report successfully generated and saved.")
                else:
                    DialogBox.errordialog("Report generation failed due to some error. "
                                          "Please close all the pdf/csv/excel files and click print again.")
        else:
            DialogBox.errordialog(errors)

    # Exchange Rate
    def getExchangeRate(self):
        month = self.from_date.dateTime().date().month()
        year = self.from_date.dateTime().date().year()

        exchg_query = "SELECT rate FROM exchange_rate WHERE month = %s AND year = %s LIMIT 1"
        value = [month, year]
        values = tuple(value, )
        exg_records = db_conn.fetch_rows(exchg_query, values)

        if exg_records:
            for row in exg_records:
                self.exg_rate = row[0]
            return True
        else:
            return False

    def showSaveDialogBox(self):
        self.print_dialog.setModal(True)

        if not self.daily_rpt_cbx.isChecked():
            self.print_pdf_cbx.setChecked(False)
            self.print_pdf_cbx.setDisabled(True)
            self.print_excel_cbx.setChecked(True)
        else:
            self.print_pdf_cbx.setChecked(True)
            self.print_pdf_cbx.setDisabled(False)
            self.print_excel_cbx.setChecked(False)

        self.print_dialog.exec_()

    def get_saveOpts(self):
        if self.print_pdf_cbx.isChecked():
            self.saveOpt.append("pdf")
        if self.print_csv_cbx.isChecked():
            self.saveOpt.append("csv")
        if self.print_excel_cbx.isChecked():
            self.saveOpt.append("excel")
        self.print_dialog.close()

    def docx_replace_regex(self, doc_obj, regex, replace):
        for p in doc_obj.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.docx_replace_regex(cell, regex, replace)

    def retranslateUi(self, ReportWindow):
        _translate = QtCore.QCoreApplication.translate
        ReportWindow.setWindowTitle(_translate("ReportWindow", "Print Report"))
        self.label_10.setText(_translate("ReportWindow", "<html><head/><body><p>"
                                                         "<img src=\"./assets/Munich_logo.png\"/>"
                                                         "</p></body></html>"))
        self.email_id.setText(_translate("ReportWindow", "Website: cgimunich.gov.in"))
        self.date_lbl.setText(_translate("ReportWindow", "Date:"))
        self.from_date.setDisplayFormat(_translate("ReportWindow", "yyyy-MM-dd"))
        self.label.setText(_translate("ReportWindow", "From:"))
        self.label_2.setText(_translate("ReportWindow", "To:"))
        self.to_date.setDisplayFormat(_translate("ReportWindow", "yyyy-MM-dd"))
        self.print_btn.setText(_translate("ReportWindow", "Print"))
        self.country_lbl.setText(_translate("ReportWindow", "Nationality:"))
        self.label_3.setText(_translate("ReportWindow", "Service:"))
        self.country_cbx.setItemText(0, _translate("ReportWindow", "Select Country"))
        self.country_cbx.setItemText(1, _translate("ReportWindow", "Other"))
        self.service_type_cbx.setItemText(0, _translate("ReportWindow", "Select Type"))
        self.bank_tr_cbx.setText(_translate("ReportWindow", "Bank Transfer"))
        self.cash_tr_cbx.setText(_translate("ReportWindow", "Cash"))
        self.card_tr_cbx.setText(_translate("ReportWindow", "Card"))
        self.gratis_sr_cbx.setText(_translate("ReportWindow", "Gratis"))
        self.rworkday_sr_cbx.setText(_translate("ReportWindow", "R Work Day"))
        self.rholiday_sr_cbx.setText(_translate("ReportWindow", "R Holiday"))
        self.postal_sr_cbx.setText(_translate("ReportWindow", "Postal Charges"))
        self.misc_sr_cbx.setText(_translate("ReportWindow", "Miscellaneous Charges"))
        self.trns_lbl.setText(_translate("ReportWindow", "Transaction Type:"))
        self.srvc_chrg_lbl.setText(_translate("ReportWindow", "Service Charges:"))
        self.print_dialog.setWindowTitle(_translate("print_dialog", "Save options"))
        self.label_4.setText(_translate("print_dialog", "Select:"))
        self.print_pdf_cbx.setText(_translate("print_dialog", "PDF"))
        self.print_csv_cbx.setText(_translate("print_dialog", "CSV"))
        self.print_excel_cbx.setText(_translate("print_dialog", "EXCEL"))
        self.print_open_btn.setText(_translate("print_dialog", "Print"))
        self.cancel_print_btn.setText(_translate("print_dialog", "Cancel"))
        self.refund_ind_cbx.setText(_translate("ReportWindow", "Refund"))
        self.exit_btn.setText(_translate("ReportWindow", "Exit"))
        self.daily_rpt_cbx.setText(_translate("ReportWindow", "Daily Report"))


if __name__ == "__main__":
    import sys

    app = QtWidgets.QApplication(sys.argv)
    ReportWindow = QtWidgets.QMainWindow()
    ui = Ui_ReportWindow()
    ui.setupUi(ReportWindow)
    ReportWindow.show()
    sys.exit(app.exec_())
